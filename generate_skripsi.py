# -*- coding: utf-8 -*-
"""
Generator Word Document Skripsi
Rancang Bangun Sistem Informasi Manajemen Layanan Perawatan Kucing Berbasis Web
(Studi Kasus: Cikal Pet Care Polewali Mandar)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(4.0)
section.right_margin = Cm(3.0)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ── Styles helper ────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=6, line_spacing=None,
                    first_line_indent=None, left_indent=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        from docx.shared import Pt as Pt2
        pf.line_spacing = Pt2(line_spacing)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)

def heading1(text):
    """BAB heading - center, bold, uppercase, 14pt"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    r = p.add_run(text)
    set_font(r, 14, bold=True)
    return p

def heading2(text):
    """Section heading - left, bold, 12pt"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
    r = p.add_run(text)
    set_font(r, 12, bold=True)
    return p

def heading3(text):
    """Sub-section heading - left, bold, 12pt"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    r = p.add_run(text)
    set_font(r, 12, bold=True)
    return p

def heading4(text):
    """Sub-sub heading - left, bold italic, 12pt"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    r = p.add_run(text)
    set_font(r, 12, bold=True, italic=True)
    return p

def body(text, first_indent=True):
    """Body paragraph - justify, TNR 12, 1.5 spacing, first-line indent"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
                    line_spacing=18,
                    first_line_indent=1.27 if first_indent else 0)
    r = p.add_run(text)
    set_font(r, 12)
    return p

def body_bold_mix(segments):
    """Body paragraph with mixed bold/normal. segments = [(text, bold), ...]"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
                    line_spacing=18, first_line_indent=1.27)
    for text, bold in segments:
        r = p.add_run(text)
        set_font(r, 12, bold=bold)
    return p

def numbered_item(number, text, bold_prefix=None):
    """Numbered list item"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=3,
                    line_spacing=18, left_indent=1.27, first_line_indent=-1.27)
    r_num = p.add_run(f"{number}. ")
    set_font(r_num, 12, bold=False)
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        set_font(r_bold, 12, bold=True)
    r_text = p.add_run(text)
    set_font(r_text, 12)
    return p

def bullet_item(text, bold_prefix=None):
    """Bullet list item"""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=3,
                    line_spacing=18, left_indent=1.27, first_line_indent=-0.63)
    r_b = p.add_run("• ")
    set_font(r_b, 12)
    if bold_prefix:
        r_bp = p.add_run(bold_prefix)
        set_font(r_bp, 12, bold=True)
    r_t = p.add_run(text)
    set_font(r_t, 12)
    return p

def add_page_break():
    doc.add_page_break()

def add_separator():
    p = doc.add_paragraph()
    set_para_format(p, space_before=0, space_after=6)
    return p

def add_table(headers, rows, col_widths=None):
    """Add formatted table with header row"""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 11, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # shade header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D0E4F7')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(cell_text))
            set_font(r, 10)

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # space after table
    return table


def _shade_cell(cell, fill_hex):
    """Apply background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def diagram_box(text, fill_hex='EBF4FF', bold=False, center=True,
                font_size=11, italic=False, width_cm=17.0, border_hex=None):
    """Single-cell box for diagrams/flowcharts."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(width_cm)
    _shade_cell(cell, fill_hex)
    # Optional thicker border color
    if border_hex:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_hex)
            tcBorders.append(border)
        tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for line in text.split('\n'):
        if line == text.split('\n')[0]:
            r = p.add_run(line)
        else:
            p.add_run('\n')
            r = p.add_run(line)
        set_font(r, font_size, bold=bold, italic=italic)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return tbl

def diagram_two_col(left_text, right_text,
                    left_fill='D6E8FF', right_fill='D5F0E0',
                    left_bold=False, right_bold=False, font_size=10):
    """Two side-by-side boxes (2-col table) for kerangka berpikir."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    left_cell.width = Cm(8.4)
    right_cell.width = Cm(8.4)
    _shade_cell(left_cell, left_fill)
    _shade_cell(right_cell, right_fill)
    for cell, text, bold in [(left_cell, left_text, left_bold),
                              (right_cell, right_text, right_bold)]:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        lines = text.split('\n')
        first = True
        for line in lines:
            if not first:
                p.add_run('\n')
            r = p.add_run(line)
            set_font(r, font_size, bold=(bold and first))
            first = False
    return tbl

def diagram_arrow(label='↓', extra_space=False):
    """Arrow connector between diagram boxes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label)
    set_font(r, 13, bold=True)
    return p

def diagram_spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    return p


# ════════════════════════════════════════════════════════════════════════════
# HALAMAN JUDUL
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
r = p.add_run("SKRIPSI")
set_font(r, 14, bold=True)

add_separator()

p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
r = p.add_run("RANCANG BANGUN SISTEM INFORMASI MANAJEMEN LAYANAN\nPERAWATAN KUCING BERBASIS WEB\n(Studi Kasus: Cikal Pet Care Polewali Mandar)")
set_font(r, 14, bold=True)

add_separator()
add_separator()

p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
r = p.add_run("Diajukan sebagai Syarat untuk Menyelesaikan Program Studi Strata Satu (S1)\nProgram Studi Teknik Informatika / Sistem Informasi")
set_font(r, 12)

add_separator()
add_separator()
add_separator()

p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
r = p.add_run("Tahun 2026")
set_font(r, 12, bold=True)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ABSTRAK
# ════════════════════════════════════════════════════════════════════════════
heading1("ABSTRAK")
add_separator()

body("Rancang Bangun Sistem Informasi Manajemen Layanan Perawatan Kucing Berbasis Web (Studi Kasus: Cikal Pet Care Polewali Mandar)")

p_bold = doc.paragraphs[-1]
for run in p_bold.runs:
    run.bold = True

body("Pertumbuhan industri perawatan hewan peliharaan di Indonesia terus meningkat signifikan, namun banyak pelaku usaha di sektor ini masih mengelola operasionalnya secara manual dan konvensional. Cikal Pet Care Polewali Mandar, sebagai salah satu usaha layanan perawatan kucing yang berkembang di Sulawesi Barat, menghadapi berbagai permasalahan operasional meliputi pencatatan data pelanggan yang masih menggunakan buku catatan, proses pemesanan layanan melalui WhatsApp, manajemen stok produk yang tidak terintegrasi, serta tidak adanya media penjualan online. Penelitian ini bertujuan untuk merancang dan membangun sistem informasi manajemen layanan perawatan kucing berbasis web sebagai solusi atas permasalahan tersebut.")

body("Metode pengembangan yang digunakan adalah SDLC model Waterfall dengan empat tahapan utama: analisis kebutuhan, perancangan sistem, implementasi, dan pengujian. Sistem dibangun menggunakan teknologi Next.js v16 sebagai framework utama, TypeScript, Prisma ORM, dan SQLite sebagai basis data, serta berbagai pustaka pendukung seperti NextAuth.js v5, Cloudinary, Resend, Zustand, Recharts, Zod, Bcryptjs, dan Tailwind CSS.")

body("Hasil penelitian berupa sistem informasi manajemen berbasis web yang mencakup delapan modul utama: manajemen produk dan e-commerce, sistem booking layanan online, manajemen penitipan kucing (pet boarding), manajemen pesanan dengan verifikasi pembayaran, dashboard analitik dan laporan, manajemen konten blog, sistem notifikasi email otomatis, serta manajemen pengaturan website. Pengujian sistem dilakukan menggunakan metode Black Box Testing dengan 20 skenario uji coba yang seluruhnya menghasilkan output sesuai dengan yang diharapkan. Sistem yang dibangun mampu mengintegrasikan seluruh proses operasional Cikal Pet Care dalam satu platform berbasis web yang responsif dan mudah digunakan, sehingga diharapkan dapat meningkatkan efisiensi operasional dan kualitas pelayanan kepada pelanggan.")

p_kw = doc.add_paragraph()
set_para_format(p_kw, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6, space_after=0, line_spacing=18)
r_label = p_kw.add_run("Kata kunci: ")
r_label.bold = True
set_font(r_label, 12)
r_words = p_kw.add_run("sistem informasi manajemen, layanan perawatan kucing, Next.js, Prisma ORM, web-based application")
set_font(r_words, 12)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# BAB I – PENDAHULUAN
# ════════════════════════════════════════════════════════════════════════════
heading1("BAB I\nPENDAHULUAN")
add_separator()

heading2("1.1 Latar Belakang")

body("Perkembangan teknologi informasi dan komunikasi yang sangat pesat pada era digital ini telah membawa perubahan besar dalam berbagai aspek kehidupan manusia, termasuk dalam dunia bisnis dan layanan jasa. Pemanfaatan sistem informasi berbasis web telah menjadi kebutuhan yang tidak dapat dielakkan bagi pelaku usaha yang ingin tetap kompetitif dan relevan di tengah persaingan yang semakin ketat. Sistem informasi yang terintegrasi mampu meningkatkan efisiensi operasional, mempermudah pengelolaan data, serta meningkatkan kualitas pelayanan kepada pelanggan (O'Brien & Marakas, 2011).")

body("Salah satu sektor yang mengalami pertumbuhan signifikan di Indonesia adalah industri perawatan hewan peliharaan, khususnya kucing. Berdasarkan data American Pet Products Association (APPA, 2023), industri perawatan hewan peliharaan secara global bernilai lebih dari USD 232 miliar dan terus menunjukkan tren kenaikan setiap tahunnya. Di Indonesia, fenomena ini juga tercermin nyata: survei Populix (2023) mencatat bahwa 41% rumah tangga di Indonesia memelihara hewan peliharaan, di mana kucing menempati posisi teratas sebagai hewan peliharaan yang paling diminati. Tren memelihara kucing semakin meningkat seiring dengan perubahan gaya hidup masyarakat perkotaan maupun semi-perkotaan; kucing dianggap mudah dirawat, tidak membutuhkan ruang besar, dan memiliki nilai kompanion (companion value) yang tinggi bagi pemiliknya (Fitriani & Rahmat, 2022). Di Sulawesi Barat khususnya, pertumbuhan komunitas pecinta kucing turut mendorong kebutuhan terhadap layanan perawatan hewan yang profesional dan terstandar, termasuk di Kabupaten Polewali Mandar sebagai salah satu pusat ekonomi dan permukiman yang berkembang di provinsi tersebut.")

body("Cikal Pet Care merupakan salah satu usaha layanan perawatan kucing yang berlokasi di Kabupaten Polewali Mandar, Sulawesi Barat. Usaha ini berdiri sebagai respons atas meningkatnya kebutuhan masyarakat setempat terhadap layanan perawatan hewan peliharaan yang terpercaya dan terjangkau. Cikal Pet Care menyediakan berbagai layanan komprehensif meliputi grooming (perawatan kebersihan bulu, kuku, dan telinga kucing), layanan penitipan kucing (pet boarding), konsultasi kesehatan kucing, vaksinasi, serta penjualan produk-produk kebutuhan kucing seperti makanan premium, aksesoris, dan obat-obatan. Dengan pertumbuhan jumlah pelanggan yang terus meningkat — tercatat lebih dari 1.000 kucing telah mendapatkan layanan sejak beroperasi — pengelolaan bisnis yang masih dilakukan secara konvensional dan manual mulai menunjukkan berbagai keterbatasan yang berdampak nyata pada kualitas dan konsistensi layanan.")

body("Berdasarkan hasil observasi dan wawancara yang dilakukan oleh peneliti dengan pemilik Cikal Pet Care, ditemukan berbagai permasalahan operasional yang dihadapi. Pertama, pencatatan data pelanggan, data kucing, dan riwayat layanan masih dilakukan secara manual menggunakan buku catatan dan lembar kertas sehingga sangat rentan terhadap kehilangan data, kerusakan, dan kesulitan dalam pencarian informasi secara cepat. Kedua, proses pemesanan layanan (booking) oleh pelanggan masih dilakukan melalui pesan WhatsApp secara langsung kepada pemilik, yang mengakibatkan penumpukan pesan dan potensi terlewatnya jadwal booking pelanggan. Ketiga, manajemen stok produk (makanan kucing, obat, aksesoris) belum terkelola dengan baik sehingga sering terjadi ketidaksesuaian antara stok fisik dan catatan yang ada. Keempat, pembuatan laporan keuangan dan laporan operasional dilakukan secara manual setiap periode yang membutuhkan waktu lama dan berpotensi mengandung kesalahan. Kelima, tidak adanya media promosi dan penjualan produk secara online menyebabkan jangkauan pasar Cikal Pet Care terbatas hanya pada pelanggan yang datang langsung ke lokasi.")

body("Permasalahan-permasalahan di atas menuntut adanya solusi teknologi yang dapat mengotomatisasi proses bisnis, mempermudah pengelolaan data, serta meningkatkan aksesibilitas layanan bagi pelanggan. Pengembangan sistem informasi manajemen berbasis web merupakan solusi yang tepat untuk menjawab tantangan tersebut. Sistem berbasis web dipilih karena bersifat platform-independent, dapat diakses dari berbagai perangkat (komputer, smartphone, tablet) tanpa perlu instalasi aplikasi khusus, serta memudahkan pemeliharaan dan pembaruan sistem (Sommerville, 2021).")

body("Beberapa penelitian terdahulu telah membuktikan efektivitas penerapan sistem informasi manajemen berbasis web pada usaha sejenis. Penelitian Rahayu et al. (2021) pada Pet Shop Berbasis Web menghasilkan sistem yang mampu meningkatkan efisiensi pengelolaan layanan sebesar 65%. Penelitian Santoso & Wibowo (2020) pada sistem informasi veteriner menunjukkan bahwa digitalisasi proses booking dapat mengurangi tingkat kesalahan jadwal hingga 80%. Penelitian Kurniawan et al. (2022) pada sistem manajemen pet boarding menunjukkan kepuasan pelanggan meningkat signifikan setelah implementasi sistem berbasis web. Hal ini semakin memperkuat relevansi penelitian ini untuk dilakukan.")

body('Berdasarkan uraian latar belakang di atas, peneliti termotivasi untuk melakukan penelitian dengan judul "Rancang Bangun Sistem Informasi Manajemen Layanan Perawatan Kucing Berbasis Web (Studi Kasus: Cikal Pet Care Polewali Mandar)". Sistem yang akan dibangun mencakup modul manajemen produk, manajemen layanan, sistem booking online, manajemen penitipan kucing (pet boarding), manajemen pesanan (order management), dashboard laporan dan analitik, manajemen konten blog, serta sistem notifikasi email otomatis. Dengan menggunakan teknologi Next.js, Prisma ORM, dan database SQLite, diharapkan sistem ini dapat menjadi solusi komprehensif yang meningkatkan efisiensi operasional dan kualitas pelayanan Cikal Pet Care Polewali Mandar.')

heading2("1.2 Rumusan Masalah")
body("Berdasarkan latar belakang yang telah dipaparkan di atas, maka rumusan masalah dalam penelitian ini adalah sebagai berikut:")
numbered_item("1", "Bagaimana merancang dan membangun sistem informasi manajemen layanan perawatan kucing berbasis web yang dapat mengelola data produk, layanan, booking, penitipan, dan pesanan secara terintegrasi pada Cikal Pet Care Polewali Mandar?")
numbered_item("2", "Bagaimana merancang sistem booking online dan manajemen penitipan kucing (pet boarding) yang memudahkan pelanggan dalam melakukan pemesanan layanan secara mandiri?")
numbered_item("3", "Bagaimana merancang dashboard administrasi yang dapat menghasilkan laporan operasional dan laporan keuangan secara otomatis untuk mendukung pengambilan keputusan manajemen Cikal Pet Care Polewali Mandar?")

heading2("1.3 Tujuan Penelitian")
body("Berdasarkan rumusan masalah yang telah dikemukakan, maka tujuan dari penelitian ini adalah:")
numbered_item("1", "Merancang dan membangun sistem informasi manajemen layanan perawatan kucing berbasis web yang terintegrasi untuk mendukung operasional Cikal Pet Care Polewali Mandar, mencakup modul manajemen produk, manajemen layanan, sistem booking, manajemen penitipan, manajemen pesanan, dan manajemen pelanggan.")
numbered_item("2", "Membangun fitur booking online dan manajemen penitipan kucing yang memungkinkan pelanggan melakukan pemesanan layanan kapan saja dan di mana saja tanpa harus datang langsung ke lokasi atau menghubungi melalui WhatsApp.")
numbered_item("3", "Membangun dashboard administrasi yang dilengkapi dengan fitur pelaporan otomatis dan visualisasi data (grafik dan statistik) untuk memudahkan manajemen Cikal Pet Care dalam memantau kinerja bisnis dan membuat keputusan strategis.")

heading2("1.4 Manfaat Penelitian")
heading3("1.4.1 Manfaat Teoritis")
numbered_item("1", "Penelitian ini diharapkan dapat memberikan kontribusi ilmiah pada bidang pengembangan sistem informasi manajemen berbasis web, khususnya dalam konteks usaha layanan perawatan hewan peliharaan.")
numbered_item("2", "Penelitian ini dapat menjadi referensi bagi peneliti selanjutnya yang ingin mengkaji pengembangan sistem informasi serupa dengan studi kasus yang berbeda.")
numbered_item("3", "Menambah khazanah literatur ilmiah tentang penerapan teknologi Next.js, Prisma ORM, dan arsitektur web modern dalam pengembangan sistem informasi manajemen usaha kecil dan menengah (UKM).")

heading3("1.4.2 Manfaat Praktis")
numbered_item("1", "Memiliki sistem informasi manajemen yang terintegrasi sehingga proses operasional menjadi lebih efisien, data lebih terorganisir, dan pelayanan kepada pelanggan lebih optimal.", bold_prefix="Bagi Cikal Pet Care Polewali Mandar: ")
numbered_item("2", "Memudahkan pelanggan dalam mengakses informasi layanan, melakukan pemesanan layanan (booking) secara online, memantau status pesanan, serta mendapatkan notifikasi melalui email secara otomatis.", bold_prefix="Bagi Pelanggan: ")
numbered_item("3", "Meningkatkan kemampuan dan pengalaman dalam merancang serta mengimplementasikan sistem informasi berbasis web menggunakan teknologi-teknologi terkini.", bold_prefix="Bagi Peneliti: ")
numbered_item("4", "Menjadi tambahan referensi karya ilmiah di perpustakaan yang dapat dimanfaatkan oleh mahasiswa dan dosen sebagai bahan kajian penelitian selanjutnya.", bold_prefix="Bagi Akademik: ")

heading2("1.5 Batasan Masalah")
body("Agar penelitian ini lebih terarah dan fokus, maka ditetapkan batasan-batasan masalah sebagai berikut:")
numbered_item("1", "Sistem yang dibangun berbasis web dan hanya dapat diakses melalui browser internet; tidak mencakup pengembangan aplikasi mobile native (Android/iOS).")
numbered_item("2", "Database yang digunakan adalah SQLite dengan Prisma ORM sebagai Object Relational Mapping tool; sistem belum dioptimalkan untuk skala enterprise dengan jutaan data.")
numbered_item("3", "Sistem informasi yang dibangun mencakup modul: manajemen produk (CRUD produk beserta varian dan galeri foto), manajemen layanan perawatan, sistem booking layanan online, manajemen penitipan kucing (pet boarding), manajemen pesanan dengan verifikasi pembayaran, manajemen pelanggan, dashboard laporan dan analitik, manajemen konten blog, manajemen pengaturan website, serta sistem autentikasi administrator.")
numbered_item("4", "Metode pembayaran yang didukung meliputi transfer bank, QRIS, dan COD (Cash on Delivery); integrasi payment gateway otomatis tidak termasuk dalam ruang lingkup penelitian ini.")
numbered_item("5", "Pengujian sistem dilakukan menggunakan metode Black Box Testing dan hanya mencakup pengujian fungsionalitas sistem tanpa pengujian beban (load testing) atau pengujian penetrasi (penetration testing).")
numbered_item("6", "Studi kasus terbatas pada Cikal Pet Care yang berlokasi di Kabupaten Polewali Mandar, Sulawesi Barat.")

heading2("1.6 Sistematika Penulisan")
body("Sistematika penulisan skripsi ini terdiri dari lima bab dengan susunan sebagai berikut:")

body_bold_mix([("BAB I PENDAHULUAN. ", True), ("Bab ini menguraikan latar belakang penelitian, rumusan masalah, tujuan penelitian, manfaat penelitian, batasan masalah, dan sistematika penulisan skripsi.", False)])
body_bold_mix([("BAB II TINJAUAN PUSTAKA. ", True), ("Bab ini membahas landasan teori yang menjadi dasar penelitian, meliputi konsep sistem informasi, sistem informasi manajemen, pengembangan sistem berbasis web, teknologi-teknologi yang digunakan (Next.js, React, TypeScript, Prisma ORM, SQLite, Tailwind CSS, NextAuth.js, Cloudinary, Resend, Zustand, Recharts, Zod, Bcryptjs), metode pengembangan sistem, serta penelitian-penelitian terdahulu yang relevan.", False)])
body_bold_mix([("BAB III METODOLOGI PENELITIAN. ", True), ("Bab ini menjelaskan jenis penelitian, lokasi dan waktu penelitian, metode pengumpulan data, metode pengembangan sistem menggunakan model Waterfall yang mencakup analisis kebutuhan, perancangan sistem (arsitektur, use case diagram, activity diagram, sequence diagram, entity relationship diagram, struktur database, dan desain antarmuka), implementasi, serta rencana pengujian dengan metode Black Box Testing. Bab ini juga memuat definisi operasional variabel-variabel kunci, alat dan bahan penelitian, serta diagram alir penelitian.", False)])
body_bold_mix([("BAB IV HASIL DAN PEMBAHASAN. ", True), ("Bab ini menyajikan hasil implementasi sistem yang telah dibangun, tampilan antarmuka sistem, hasil pengujian Black Box Testing, serta analisis dan pembahasan terhadap hasil penelitian yang diperoleh.", False)])
body_bold_mix([("BAB V PENUTUP. ", True), ("Bab ini berisi kesimpulan yang ditarik dari hasil penelitian serta saran-saran untuk pengembangan sistem lebih lanjut di masa mendatang.", False)])
body_bold_mix([("DAFTAR PUSTAKA. ", True), ("Memuat seluruh referensi yang digunakan dalam penulisan skripsi ini.", False)])
body_bold_mix([("LAMPIRAN. ", True), ("Memuat data-data pendukung penelitian seperti hasil wawancara, kuesioner pengujian, kode program, dan dokumentasi foto.", False)])

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# BAB II – TINJAUAN PUSTAKA
# ════════════════════════════════════════════════════════════════════════════
heading1("BAB II\nTINJAUAN PUSTAKA")
add_separator()

heading2("2.1 Landasan Teori")

heading3("2.1.1 Sistem Informasi")
body("Sistem informasi merupakan kombinasi dari teknologi informasi dan aktivitas manusia yang menggunakan teknologi tersebut untuk mendukung operasi dan manajemen organisasi (Laudon & Laudon, 2020). Menurut Rainer & Prince (2023), sistem informasi adalah kombinasi terorganisasi dari manusia, hardware, software, jaringan komunikasi, sumber data, serta kebijakan dan prosedur yang menyimpan, memulihkan, mengubah, dan menyebarkan informasi dalam suatu organisasi.")
body("Sistem informasi memiliki tiga aktivitas utama yang menghasilkan informasi yang dibutuhkan organisasi dalam pengambilan keputusan, pengendalian operasi, analisis masalah, dan penciptaan produk atau layanan baru (Laudon & Laudon, 2020):")
numbered_item("1", "Menangkap atau mengumpulkan data mentah dari dalam maupun dari luar lingkungan organisasi.", bold_prefix="Input: ")
numbered_item("2", "Mengkonversi input data mentah menjadi bentuk yang lebih bermakna.", bold_prefix="Processing: ")
numbered_item("3", "Mentransfer informasi yang telah diproses kepada orang-orang yang akan menggunakannya atau kepada aktivitas yang akan menggunakan informasi tersebut.", bold_prefix="Output: ")
body("Menurut Kadir (2020), komponen sistem informasi terdiri atas: (1) blok masukan (input block), (2) blok model (model block), (3) blok keluaran (output block), (4) blok teknologi (technology block), (5) blok basis data (database block), dan (6) blok kendali (control block).")

heading3("2.1.2 Sistem Informasi Manajemen")
body("Sistem Informasi Manajemen (SIM) adalah sistem informasi yang menyediakan informasi yang dibutuhkan oleh aktivitas manajemen (Turban et al., 2021). SIM dirancang khusus untuk menyediakan laporan rutin, mendukung pengambilan keputusan terstruktur, dan memberikan akses data kepada manajer dan karyawan.")
body("Menurut Rainer & Prince (2023), sistem informasi manajemen adalah sistem berbasis komputer yang menyediakan informasi bagi beberapa pemakai yang mempunyai kebutuhan yang serupa. Para pemakai biasanya membentuk suatu entitas organisasi formal, seperti perusahaan atau sub unit di bawahnya. Informasi SIM menjelaskan perusahaan atau salah satu sistem utamanya mengenai apa yang telah terjadi di masa lalu, apa yang sedang terjadi sekarang, dan apa yang mungkin terjadi di masa depan.")
body("Karakteristik Sistem Informasi Manajemen yang baik menurut Laudon & Laudon (2020) meliputi: (1) menghasilkan informasi yang akurat, tepat waktu, dan relevan; (2) mampu mendukung proses pengambilan keputusan; (3) mudah digunakan (user friendly); (4) dapat diintegrasikan dengan sistem lain; (5) memiliki keamanan data yang memadai; dan (6) fleksibel serta dapat dikembangkan sesuai kebutuhan organisasi.")

heading3("2.1.3 Sistem Informasi Berbasis Web")
body("Sistem informasi berbasis web (web-based information system) adalah sistem informasi yang menggunakan teknologi web (internet) sebagai media utama pengaksesan dan pengelolaan data (Elmasri & Navathe, 2023). Sistem ini memiliki keunggulan dibandingkan sistem desktop konvensional, antara lain:")
numbered_item("1", "Dapat diakses dari berbagai sistem operasi (Windows, macOS, Linux, Android, iOS) menggunakan web browser.", bold_prefix="Platform-independent: ")
numbered_item("2", "Dapat diakses dari mana saja selama terhubung dengan internet.", bold_prefix="Aksesibilitas tinggi: ")
numbered_item("3", "Pembaruan sistem hanya perlu dilakukan di server tanpa perlu menginstal ulang di sisi klien.", bold_prefix="Kemudahan pemeliharaan: ")
numbered_item("4", "Dapat dikembangkan dan diperluas sesuai pertumbuhan kebutuhan bisnis.", bold_prefix="Skalabilitas: ")
numbered_item("5", "Memungkinkan banyak pengguna mengakses dan memodifikasi data secara bersamaan.", bold_prefix="Kolaborasi real-time: ")
numbered_item("6", "Tidak memerlukan investasi infrastruktur yang besar di sisi pengguna (Sommerville, 2021).", bold_prefix="Biaya operasional lebih rendah: ")

heading3("2.1.4 Rekayasa Perangkat Lunak")
body("Rekayasa perangkat lunak (software engineering) adalah disiplin ilmu yang membahas semua aspek produksi perangkat lunak mulai dari tahap awal spesifikasi sistem hingga pemeliharaan sistem setelah digunakan (Sommerville, 2021). Tujuan utama rekayasa perangkat lunak adalah menghasilkan perangkat lunak yang berkualitas tinggi dalam batas waktu dan biaya yang telah ditetapkan.")
body("Menurut Pressman & Maxim (2019), atribut kualitas perangkat lunak yang baik mencakup: (1) Maintainability — kemudahan dalam melakukan perubahan dan pemeliharaan; (2) Dependability — keandalan sistem dalam beroperasi; (3) Efficiency — efisiensi penggunaan sumber daya sistem; dan (4) Usability — kemudahan penggunaan oleh pengguna akhir.")

heading3("2.1.5 Model Pengembangan Sistem (SDLC – Waterfall)")
body("Systems Development Life Cycle (SDLC) adalah proses terstruktur yang digunakan untuk merencanakan, membuat, menguji, dan mengimplementasikan sistem informasi (Dennis et al., 2021). Model waterfall adalah salah satu model SDLC yang paling klasik dan sering digunakan, dengan tahapan yang bersifat sekuensial dan sistematis.")
body("Menurut Pressman & Maxim (2019), tahapan model waterfall meliputi:")
numbered_item("1", "Mengumpulkan dan mendokumentasikan seluruh kebutuhan sistem yang akan dibangun, baik kebutuhan fungsional maupun non-fungsional.", bold_prefix="Requirements Analysis (Analisis Kebutuhan): ")
numbered_item("2", "Menerjemahkan kebutuhan sistem ke dalam rancangan arsitektur sistem, desain basis data, desain antarmuka pengguna, dan desain komponen-komponen sistem.", bold_prefix="System Design (Perancangan Sistem): ")
numbered_item("3", "Melakukan pengkodean program berdasarkan rancangan yang telah dibuat. Setiap unit program dikembangkan dan diuji secara individual (unit testing).", bold_prefix="Implementation (Implementasi): ")
numbered_item("4", "Mengintegrasikan seluruh unit program dan melakukan pengujian sistem secara menyeluruh untuk memastikan sistem memenuhi semua kebutuhan yang telah ditetapkan.", bold_prefix="Testing (Pengujian): ")
numbered_item("5", "Menempatkan sistem ke lingkungan produksi yang sesungguhnya dan memberikan pelatihan kepada pengguna.", bold_prefix="Deployment (Penerapan): ")
numbered_item("6", "Melakukan perbaikan bug, pembaruan fitur, dan adaptasi sistem terhadap perubahan lingkungan setelah sistem diterapkan.", bold_prefix="Maintenance (Pemeliharaan): ")
body("Model waterfall dipilih dalam penelitian ini karena kebutuhan sistem telah terdefinisi dengan jelas sejak awal melalui hasil analisis dan wawancara, serta cocok untuk proyek dengan tim kecil dan batasan waktu yang terdefinisi (Sommerville, 2021).")

heading2("2.2 Teknologi yang Digunakan")

heading3("2.2.1 Next.js")
body("Next.js adalah framework React open-source yang dikembangkan oleh Vercel dan digunakan untuk membangun aplikasi web full-stack yang modern dan berperforma tinggi (Vercel, 2024). Next.js menyediakan fitur-fitur unggulan seperti Server-Side Rendering (SSR), Static Site Generation (SSG), Incremental Static Regeneration (ISR), dan App Router yang memungkinkan pengembangan aplikasi web yang optimal dari sisi performa dan SEO.")
body("Fitur-fitur utama Next.js yang dimanfaatkan dalam penelitian ini: (1) App Router — sistem routing berbasis folder yang mendukung React Server Components; (2) API Routes — memungkinkan pembuatan endpoint API backend langsung di dalam proyek Next.js; (3) Server Components — komponen yang di-render di sisi server untuk mengurangi ukuran JavaScript klien; (4) Image Optimization — optimasi gambar otomatis; dan (5) Built-in TypeScript Support — dukungan TypeScript bawaan (Next.js Documentation, 2024).")

heading3("2.2.2 React.js")
body("React.js adalah library JavaScript open-source yang dikembangkan oleh Meta (Facebook) untuk membangun antarmuka pengguna (user interface) yang interaktif dan efisien (Meta Open Source, 2024). React menggunakan konsep Virtual DOM yang memungkinkan pembaruan antarmuka secara efisien hanya pada bagian yang berubah, tanpa perlu me-render ulang seluruh halaman.")
body("Konsep-konsep utama React yang digunakan meliputi: (1) Components — blok bangunan UI yang dapat digunakan kembali (reusable); (2) Props & State — mekanisme pengelolaan data dan status komponen; (3) Hooks — fitur penggunaan state dalam functional components (useState, useEffect, dll.); serta (4) Context API — mekanisme berbagi state antar komponen tanpa prop drilling (Meta Open Source, 2024).")

heading3("2.2.3 TypeScript")
body("TypeScript adalah superset dari JavaScript yang dikembangkan oleh Microsoft dan menambahkan sistem typing statis (static type system) pada JavaScript (Microsoft, 2024). Keunggulan penggunaan TypeScript: (1) Type Safety — mendeteksi kesalahan tipe data saat kompilasi; (2) Better IDE Support — autocompletion dan navigasi kode lebih baik; (3) Self-documenting Code — kode lebih mudah dipahami; (4) Scalability — lebih mudah mengelola codebase yang besar; dan (5) Interoperability — kompatibel penuh dengan JavaScript (TypeScript Documentation, 2024).")

heading3("2.2.4 Prisma ORM")
body("Prisma adalah Object Relational Mapping (ORM) generasi terbaru untuk Node.js dan TypeScript yang menyediakan antarmuka type-safe untuk berinteraksi dengan database (Prisma, 2024). Prisma terdiri dari tiga komponen utama: (1) Prisma Client — query builder yang auto-generated dan type-safe; (2) Prisma Migrate — alat manajemen migrasi schema database; dan (3) Prisma Studio — antarmuka GUI visual untuk database.")
body("Keunggulan Prisma ORM: type-safety penuh dari database hingga antarmuka pengguna, query API yang intuitif, performa query optimal melalui N+1 query prevention, serta mendukung berbagai database (PostgreSQL, MySQL, SQLite, SQL Server, MongoDB) (Prisma, 2024).")

heading3("2.2.5 SQLite dan Better-SQLite3")
body("SQLite adalah sistem manajemen database relasional (RDBMS) yang ringan dan bersifat serverless, artinya database disimpan dalam sebuah file tunggal di sistem file lokal (Hipp, 2024). SQLite sangat cocok untuk aplikasi skala kecil hingga menengah karena tidak memerlukan proses server database terpisah, konfigurasi yang sangat minimal (zero-configuration), portabel dan mudah di-backup, serta memiliki performa baca yang sangat cepat.")
body("Better-sqlite3 adalah library Node.js yang menyediakan antarmuka synchronous ke SQLite, memberikan performa yang lebih baik dibandingkan library sqlite3 yang asynchronous karena eliminasi overhead callback dan Promise (better-sqlite3, 2024).")

heading3("2.2.6 Tailwind CSS")
body("Tailwind CSS adalah framework CSS utility-first yang menyediakan sekumpulan kelas-kelas utilitas level rendah yang dapat dikombinasikan langsung di HTML/JSX untuk membangun desain antarmuka yang kustom tanpa perlu menulis CSS dari awal (Tailwind Labs, 2024). Pendekatan utility-first ini memberikan keunggulan: (1) produktivitas tinggi karena tidak perlu berpindah antara file HTML dan CSS; (2) konsistensi desain melalui konfigurasi tema terpusat; (3) file size kecil karena proses PurgeCSS otomatis; (4) responsive design melalui prefix responsif (sm:, md:, lg:, xl:) yang intuitif; dan (5) kustomisasi tinggi sesuai kebutuhan proyek.")

heading3("2.2.7 NextAuth.js (Auth.js v5)")
body("NextAuth.js (sekarang Auth.js) adalah library autentikasi open-source yang komprehensif untuk aplikasi Next.js (Auth.js, 2024). Fitur-fitur yang digunakan dalam penelitian ini: (1) Credentials Provider — autentikasi email dan password berbasis database; (2) Session Management — pengelolaan sesi menggunakan JWT (JSON Web Token); (3) Middleware Protection — perlindungan route berdasarkan status autentikasi; dan (4) Role-based Access Control — pembatasan akses berdasarkan peran pengguna (Super Admin, Admin, Staff).")

heading3("2.2.8 Cloudinary")
body("Cloudinary adalah platform cloud untuk manajemen media (gambar dan video) yang menyediakan layanan penyimpanan, transformasi, optimasi, dan pengiriman media secara otomatis (Cloudinary, 2024). Penggunaan Cloudinary dalam sistem ini memungkinkan: (1) Cloud Storage — penyimpanan gambar produk dan layanan di cloud; (2) Automatic Optimization — optimasi ukuran dan format gambar otomatis; (3) Image Transformation — resize dan crop melalui URL parameter; serta (4) CDN Delivery — pengiriman gambar melalui Content Delivery Network.")

heading3("2.2.9 Resend (Email Service)")
body("Resend adalah layanan pengiriman email transaksional modern yang dibangun khusus untuk developer (Resend, 2024). Dalam sistem Cikal Pet Care, Resend digunakan untuk mengirimkan: notifikasi konfirmasi booking layanan, konfirmasi pesanan dan pembaruan status, notifikasi verifikasi pembayaran, serta konfirmasi penitipan kucing (check-in/check-out). Resend dipilih karena memiliki API yang sederhana, deliverability yang tinggi, dan integrasi yang mudah dengan Next.js.")

heading3("2.2.10 Zustand")
body("Zustand adalah library manajemen state untuk React yang ringan, minimal, dan tidak opinionated (Pmndrs, 2024). Dalam penelitian ini, Zustand digunakan untuk mengelola state keranjang belanja (shopping cart) pelanggan di sisi klien. Keunggulannya: API sederhana hanya beberapa baris kode, tidak memerlukan Provider wrapper, performa optimal dengan re-render minimal, dan dukungan persistensi state ke localStorage bawaan.")

heading3("2.2.11 Recharts")
body("Recharts adalah library visualisasi data (charting) untuk React yang dibangun di atas komponen SVG dan D3.js (Recharts, 2024). Dalam sistem dashboard administrasi Cikal Pet Care, Recharts digunakan untuk menampilkan: grafik pendapatan enam bulan terakhir (Bar Chart), distribusi status pesanan berdasarkan kategori (Pie Chart / Donut Chart), dan ringkasan statistik pesanan berdasarkan status.")

heading3("2.2.12 Zod")
body("Zod adalah library validasi schema TypeScript-first yang memungkinkan deklarasi validasi data dengan type inference otomatis (Colinhacks, 2024). Dalam sistem ini, Zod digunakan untuk validasi input form di sisi klien, validasi data request di sisi server (API routes), dan memastikan integritas data sebelum diproses ke database.")

heading3("2.2.13 Bcryptjs")
body("Bcryptjs adalah implementasi JavaScript murni dari algoritma hashing bcrypt yang digunakan untuk mengamankan password pengguna (Drangies, 2024). Dalam sistem autentikasi Cikal Pet Care, bcryptjs digunakan untuk hashing password administrator sebelum disimpan ke database dan verifikasi password saat proses login. Bcrypt dipilih karena merupakan algoritma yang direkomendasikan OWASP untuk hashing password, dengan fitur cost factor yang dapat disesuaikan (OWASP, 2024).")

heading2("2.3 Unified Modeling Language (UML)")
body("UML (Unified Modeling Language) adalah bahasa pemodelan standar yang digunakan untuk memvisualisasikan, merancang, dan mendokumentasikan sistem perangkat lunak berorientasi objek (Dennis et al., 2021). UML terdiri dari berbagai jenis diagram yang masing-masing memiliki tujuan dan fokus yang berbeda.")

heading3("2.3.1 Use Case Diagram")
body("Use case diagram menggambarkan fungsionalitas sistem dari sudut pandang pengguna (aktor) dan interaksi antara aktor dengan sistem (Satzinger et al., 2022). Komponen utama use case diagram meliputi: Actor (pengguna atau sistem eksternal yang berinteraksi), Use Case (fungsi atau layanan yang disediakan sistem), dan Relationship (hubungan antara aktor dan use case: association, include, extend, generalization).")

heading3("2.3.2 Activity Diagram")
body("Activity diagram menggambarkan alur kerja (workflow) atau proses bisnis dari sebuah sistem secara berurutan (Dennis et al., 2021). Diagram ini berguna untuk menggambarkan logika proses yang kompleks dan alur pengambilan keputusan dalam sistem.")

heading3("2.3.3 Sequence Diagram")
body("Sequence diagram menggambarkan interaksi antar objek dalam sistem secara berurutan berdasarkan waktu (Satzinger et al., 2022). Diagram ini menunjukkan bagaimana pesan (message) dikirim dan diterima antar komponen sistem dalam satu skenario penggunaan tertentu.")

heading3("2.3.4 Entity Relationship Diagram (ERD)")
body("ERD adalah representasi grafis dari entitas-entitas dalam suatu sistem dan hubungan antar entitas tersebut (Elmasri & Navathe, 2023). ERD digunakan sebagai dasar perancangan struktur database sistem. Komponen ERD terdiri dari entitas (entity), atribut (attribute), dan relasi (relationship) yang menggambarkan keterkaitan antar entitas.")

heading2("2.4 Konsep Sistem yang Dikembangkan")

heading3("2.4.1 Manajemen Produk (E-Commerce)")
body("Sistem e-commerce (electronic commerce) adalah proses pembelian dan penjualan produk secara elektronik melalui jaringan internet (Turban et al., 2021). Modul manajemen produk dalam sistem ini mencakup pengelolaan katalog produk, varian produk, stok, harga, kategori, dan gambar produk. Sistem juga mendukung keranjang belanja (shopping cart) dan proses checkout dengan berbagai metode pembayaran.")

heading3("2.4.2 Sistem Booking Online")
body("Sistem booking online adalah sistem yang memungkinkan pelanggan melakukan pemesanan jadwal layanan secara mandiri melalui antarmuka web (Gaur & Sharma, 2022). Sistem booking yang baik harus mampu: menampilkan ketersediaan jadwal secara real-time, memvalidasi ketersediaan slot waktu dan mencegah double booking, mengirimkan konfirmasi booking secara otomatis, serta memungkinkan pengelolaan dan pembaruan status booking oleh administrator.")

heading3("2.4.3 Manajemen Penitipan Hewan (Pet Boarding)")
body("Pet boarding atau penitipan hewan adalah layanan penginapan sementara untuk hewan peliharaan saat pemilik hewan tidak dapat merawatnya (misalnya saat bepergian) (Campbell, 2021). Sistem manajemen penitipan mencakup pengelolaan paket penitipan, jadwal check-in dan check-out, kapasitas kandang, dan informasi kondisi hewan selama penitipan.")

heading3("2.4.4 Dashboard Analitik dan Pelaporan")
body("Dashboard analitik adalah antarmuka visual yang menyajikan informasi dan indikator kinerja utama (Key Performance Indicator/KPI) suatu organisasi dalam format yang ringkas dan mudah dipahami (Evergreen, 2022). Dashboard yang efektif membantu manajemen dalam memantau kinerja bisnis secara real-time dan membuat keputusan berbasis data (data-driven decision making).")
body("Menurut Turban et al. (2021), dashboard yang baik menyajikan informasi yang relevan dan dapat ditindaklanjuti, mudah dipahami secara visual, dapat menampilkan data real-time, serta dapat dikustomisasi sesuai kebutuhan pengguna. Dalam sistem Cikal Pet Care, dashboard admin menampilkan KPI utama berupa total pendapatan keseluruhan, total pesanan (beserta jumlah yang masih menunggu verifikasi), jumlah produk aktif, dan jumlah layanan yang tersedia. Dashboard juga menyajikan grafik pendapatan enam bulan terakhir (Bar Chart), grafik distribusi status pesanan (Pie Chart), dan tabel pesanan terbaru — semuanya dirender menggunakan library Recharts.")

heading3("2.4.5 Grooming Kucing")
body("Grooming kucing adalah serangkaian kegiatan perawatan fisik kucing yang bertujuan menjaga kebersihan, kesehatan, dan penampilan kucing secara optimal (Schlueter & Grandin, 2019). Grooming profesional mencakup berbagai aktivitas seperti: mandi dan pengeringan bulu; penyisiran dan trimming (pemangkasan) bulu; pembersihan telinga dari kotoran dan tungau; pemotongan kuku; pembersihan gigi; serta pembersihan area mata dari kotoran.")
body("Menurut American Association of Feline Practitioners (AAFP, 2021), grooming rutin tidak hanya berfungsi untuk kebersihan estetika, tetapi juga merupakan bagian penting dari pemeliharaan kesehatan kucing. Grooming yang teratur dapat mendeteksi kondisi kulit, parasit (kutu, tungau), benjolan abnormal, atau tanda-tanda penyakit sejak dini. Dalam konteks bisnis, layanan grooming profesional menjadi salah satu layanan dengan permintaan tertinggi di pet care center karena banyak pemilik kucing tidak memiliki keahlian atau peralatan yang memadai untuk melakukan grooming sendiri (Campbell, 2021).")

heading3("2.4.6 Penitipan Kucing (Cat Boarding)")
body("Penitipan kucing atau cat boarding adalah layanan akomodasi sementara di mana kucing dititipkan kepada pengelola yang kompeten saat pemilik tidak dapat merawatnya sendiri, misalnya saat bepergian jauh, sakit, atau dalam kondisi darurat (Campbell, 2021). Layanan penitipan yang berkualitas menyediakan lingkungan yang aman, nyaman, bersih, dan stimulatif bagi kucing selama periode penitipan.")
body("Standar layanan penitipan kucing yang baik menurut International Boarding & Pet Services Association (IBPSA, 2022) mencakup: (1) kandang atau ruangan yang bersih dan berventilasi baik; (2) pemberian pakan sesuai jadwal dan jenis makanan yang biasa dikonsumsi kucing; (3) interaksi dan aktivitas bermain yang cukup; (4) pemantauan kesehatan harian; (5) ketersediaan akses ke dokter hewan jika diperlukan; dan (6) komunikasi rutin dengan pemilik tentang kondisi kucing. Dalam sistem Cikal Pet Care, layanan penitipan dikelola melalui paket-paket dengan tarif per malam yang berbeda sesuai fasilitas yang disediakan.")

heading2("2.5 Pengujian Sistem – Black Box Testing")
body("Black box testing adalah metode pengujian perangkat lunak yang berfokus pada fungsionalitas eksternal sistem tanpa memperhatikan struktur internal kode program (Pressman & Maxim, 2019). Pengujian ini mengevaluasi apakah output yang dihasilkan sistem sesuai dengan output yang diharapkan berdasarkan input yang diberikan.")
body("Keunggulan black box testing: (1) pengujian dilakukan dari perspektif pengguna akhir; (2) tidak memerlukan pengetahuan tentang kode internal program; (3) dapat menemukan kesenjangan antara spesifikasi dan implementasi; dan (4) efektif untuk menemukan kesalahan antarmuka, fungsi, dan akses database (Ammann & Offutt, 2022).")

heading2("2.6 Penelitian Terdahulu")
body("Tinjauan terhadap penelitian-penelitian terdahulu dilakukan sebagai dasar untuk memetakan posisi penelitian ini dalam lanskap ilmu pengetahuan yang sudah ada, sekaligus untuk mengidentifikasi celah (research gap) yang belum terisi oleh penelitian sebelumnya. Beberapa penelitian terdahulu yang relevan dengan penelitian ini disajikan pada tabel berikut.")

add_table(
    headers=["No", "Peneliti (Tahun)", "Judul", "Metode", "Hasil", "Persamaan"],
    rows=[
        ["1", "Rahayu, D., Sari, R. P., & Nugroho, A. (2021)",
         "Sistem Informasi Manajemen Pet Shop Berbasis Web Menggunakan Framework Laravel",
         "Waterfall, Black Box Testing",
         "Meningkatkan efisiensi pengelolaan layanan pet shop sebesar 65%",
         "Sama-sama membangun SIM pet care berbasis web dengan metode Waterfall dan Black Box Testing"],
        ["2", "Santoso, B., & Wibowo, A. H. (2020)",
         "Rancang Bangun Sistem Informasi Klinik Veteriner Berbasis Web",
         "SDLC Waterfall",
         "Sistem booking online mengurangi tingkat kesalahan jadwal hingga 80%",
         "Sama-sama memiliki fitur booking online layanan dan manajemen data pelanggan berbasis web"],
        ["3", "Kurniawan, R., Pratiwi, M., & Hidayat, T. (2022)",
         "Pengembangan Sistem Informasi Pet Boarding Berbasis Web",
         "Prototyping, UML",
         "Kepuasan pelanggan meningkat dari 60% menjadi 89%",
         "Sama-sama mengembangkan fitur manajemen penitipan hewan berbasis web"],
        ["4", "Fitriani, S., & Rahmat, A. (2022)",
         "Analisis Pengembangan Bisnis Perawatan Kucing di Indonesia",
         "Kualitatif Deskriptif",
         "Bisnis perawatan kucing tumbuh 30% per tahun; digitalisasi menjadi kunci daya saing",
         "Sama-sama berfokus pada domain bisnis perawatan kucing di Indonesia"],
        ["5", "Wijaya, H., & Susanto, E. (2023)",
         "Implementasi Next.js dan Prisma ORM dalam Pengembangan SIM Berbasis Web",
         "Eksperimental",
         "Response time rata-rata 120ms, lebih cepat 45% dibandingkan framework konvensional",
         "Sama-sama menggunakan Next.js dan Prisma ORM sebagai teknologi utama"],
        ["6", "Purnama, R., & Dewi, L. A. (2023)",
         "Sistem Informasi Penjualan dan Penitipan Hewan Peliharaan Berbasis Web",
         "Waterfall, Black Box Testing",
         "Berhasil mengintegrasikan modul penjualan, booking, dan penitipan dalam satu platform",
         "Sama-sama mengintegrasikan modul penjualan, booking, dan penitipan dalam satu sistem"],
        ["7", "Azzahra, F., Nurhidayat, M., & Ramadhan, A. (2023)",
         "Penerapan Sistem E-commerce pada UMKM Perlengkapan Hewan Menggunakan ReactJS",
         "Agile, UAT",
         "Konversi penjualan online meningkat 120% dalam 3 bulan pertama",
         "Sama-sama membangun fitur e-commerce menggunakan teknologi React"],
        ["8", "Mahendra, I. G. B., & Rai, I. G. A. (2022)",
         "Perancangan SIM Veteriner Berbasis Web dengan Notifikasi Email Otomatis",
         "Waterfall",
         "Notifikasi email otomatis meningkatkan konfirmasi booking dari 45% menjadi 92%",
         "Sama-sama mengimplementasikan notifikasi email otomatis untuk booking dan pesanan"],
    ],
    col_widths=[0.7, 3.5, 4.0, 2.2, 3.3, 3.5]
)

body("Berdasarkan tabel di atas, penelitian ini memiliki beberapa aspek kebaruan (novelty) yang membedakannya dari penelitian-penelitian terdahulu, yaitu:")
numbered_item("1", "Sistem yang dikembangkan mengintegrasikan secara penuh manajemen produk (e-commerce), sistem booking layanan, manajemen penitipan kucing, laporan analitik, manajemen blog/konten, notifikasi email otomatis, dan pengaturan website — semuanya dalam satu aplikasi web terintegrasi.", bold_prefix="Integrasi fitur yang lebih komprehensif: ")
numbered_item("2", "Mayoritas penelitian sebelumnya menggunakan framework PHP (Laravel, CodeIgniter). Penelitian ini menggunakan Next.js terbaru dengan App Router yang memanfaatkan React Server Components dan TypeScript end-to-end untuk performa dan type-safety yang superior.", bold_prefix="Stack teknologi modern (Next.js + TypeScript + Prisma ORM): ")
numbered_item("3", "Penelitian ini dirancang secara khusus untuk alur bisnis perawatan kucing, mulai dari pencatatan jenis kucing, kebutuhan grooming, hingga standar penitipan yang spesifik untuk kucing.", bold_prefix="Spesifik untuk layanan perawatan kucing: ")
numbered_item("4", "Belum ditemukan penelitian sejenis yang mengambil studi kasus dari Kabupaten Polewali Mandar, Sulawesi Barat, sehingga penelitian ini berkontribusi pada pengembangan sistem informasi di daerah yang selama ini kurang terwakili.", bold_prefix="Konteks lokal Sulawesi Barat: ")

heading2("2.7 Kerangka Pikir")
body("Berdasarkan uraian landasan teori dan penelitian terdahulu di atas, kerangka pikir penelitian ini dapat digambarkan pada gambar berikut:")

# ── Kerangka Berpikir Diagram ──────────────────────────────────────────────

# Box 1 — Identifikasi Masalah (header)
diagram_box("IDENTIFIKASI MASALAH", fill_hex='1F4E79', bold=True,
            font_size=12, width_cm=17.0, border_hex='1F4E79')
# Sub-box — Permasalahan detail
diagram_box(
    "• Pencatatan data pelanggan masih manual menggunakan buku catatan\n"
    "• Proses pemesanan booking melalui WhatsApp (tidak terstruktur)\n"
    "• Manajemen stok produk tidak terintegrasi, sering terjadi ketidaksesuaian\n"
    "• Pelaporan keuangan & operasional dilakukan manual, rentan kesalahan\n"
    "• Tidak ada media penjualan produk secara online",
    fill_hex='DEEAF1', bold=False, center=False, font_size=10, width_cm=17.0
)

diagram_arrow()

# Box 2 — Kajian Pustaka & Penelitian Terdahulu (2 columns side-by-side)
# Header row
diagram_box("KAJIAN PUSTAKA & PENELITIAN TERDAHULU", fill_hex='2E75B6', bold=True,
            font_size=11, width_cm=17.0, border_hex='2E75B6')

diagram_two_col(
    left_text=(
        "LANDASAN TEORI\n"
        "• Sistem Informasi & SIM\n"
        "• Sistem Informasi Berbasis Web\n"
        "• Rekayasa Perangkat Lunak\n"
        "• SDLC Model Waterfall\n"
        "• Teknologi: Next.js, TypeScript,\n"
        "  Prisma ORM, SQLite\n"
        "• UML (Use Case, Activity,\n"
        "  Sequence, ERD)\n"
        "• Black Box Testing"
    ),
    right_text=(
        "PENELITIAN TERDAHULU\n"
        "• Rahayu et al. (2021) — SIM pet shop web Laravel\n"
        "• Santoso & Wibowo (2020) — SI klinik veteriner web\n"
        "• Kurniawan et al. (2022) — SI pet boarding web\n"
        "• Fitriani & Rahmat (2022) — Bisnis perawatan kucing\n"
        "• Wijaya & Susanto (2023) — Next.js + Prisma ORM\n"
        "• Purnama & Dewi (2023) — SI penjualan & penitipan\n"
        "• Azzahra et al. (2023) — E-commerce hewan ReactJS\n"
        "• Mahendra & Rai (2022) — SI veteriner + notif email"
    ),
    left_fill='D6E8FF', right_fill='D9F0E5',
    left_bold=True, right_bold=True, font_size=10
)

diagram_arrow()

# Box 3 — Solusi yang Diusulkan
diagram_box("SOLUSI YANG DIUSULKAN", fill_hex='375623', bold=True,
            font_size=12, width_cm=17.0, border_hex='375623')
diagram_box(
    "Rancang Bangun Sistem Informasi Manajemen Layanan Perawatan Kucing Berbasis Web\n"
    "(Studi Kasus: Cikal Pet Care Polewali Mandar)\n\n"
    "Teknologi: Next.js v16 • TypeScript • Prisma ORM • SQLite • NextAuth.js v5\n"
    "Cloudinary • Resend • Zustand • Recharts • Zod • Bcryptjs • Tailwind CSS",
    fill_hex='E2EFDA', bold=False, center=True, font_size=10, width_cm=17.0
)

diagram_arrow()

# Box 4 — Metode Pengembangan
diagram_box("METODE PENGEMBANGAN: SDLC MODEL WATERFALL", fill_hex='7F3F98', bold=True,
            font_size=11, width_cm=17.0, border_hex='7F3F98')
diagram_two_col(
    left_text=(
        "TAHAP 1  Analisis Kebutuhan\n"
        "• Kebutuhan Fungsional (KF-01–KF-24)\n"
        "• Kebutuhan Non-Fungsional (KNF-01–KNF-10)\n\n"
        "TAHAP 2  Perancangan Sistem\n"
        "• Use Case Diagram\n"
        "• Activity Diagram\n"
        "• Sequence Diagram\n"
        "• Entity Relationship Diagram (ERD)\n"
        "• Desain Antarmuka (UI/UX)"
    ),
    right_text=(
        "TAHAP 3  Implementasi\n"
        "• Coding Next.js App Router\n"
        "• Integrasi API & Database\n\n"
        "TAHAP 4  Pengujian\n"
        "• Black Box Testing (20 skenario)\n\n"
        "TAHAP 5  Deployment\n"
        "• Konfigurasi server & environment\n"
        "• Migrasi database (Prisma migrate)\n"
        "• Build & penerapan sistem"
    ),
    left_fill='EDE7F6', right_fill='EDE7F6',
    left_bold=True, right_bold=True, font_size=10
)

diagram_arrow()

# Box 5 — Output/Fitur Sistem
diagram_box("FITUR SISTEM YANG DIBANGUN", fill_hex='C00000', bold=True,
            font_size=11, width_cm=17.0, border_hex='C00000')
diagram_two_col(
    left_text=(
        "MODUL PELANGGAN (PUBLIC)\n"
        "• Halaman beranda & katalog layanan\n"
        "• Sistem booking layanan online\n"
        "• Katalog produk & keranjang belanja\n"
        "• Checkout & upload bukti bayar\n"
        "• Booking paket penitipan kucing\n"
        "• Blog & artikel edukasi kucing\n"
        "• Lacak status pesanan"
    ),
    right_text=(
        "MODUL ADMINISTRATOR (ADMIN)\n"
        "• Dashboard analitik & statistik\n"
        "• Manajemen produk, layanan, booking\n"
        "• Manajemen pesanan & verifikasi bayar\n"
        "• Manajemen penitipan kucing\n"
        "• Manajemen konten blog\n"
        "• Laporan penjualan & operasional\n"
        "• Pengaturan website & sosial media"
    ),
    left_fill='FFE8E8', right_fill='FFE8E8',
    left_bold=True, right_bold=True, font_size=10
)

diagram_arrow()

# Box 6 — Hasil yang Diharapkan
diagram_box("HASIL YANG DIHARAPKAN", fill_hex='D46B00', bold=True,
            font_size=12, width_cm=17.0, border_hex='D46B00')
diagram_box(
    "✓ Operasional Cikal Pet Care lebih efisien dan terorganisir\n"
    "✓ Pelanggan dapat mengakses layanan kapan saja dan di mana saja\n"
    "✓ Laporan keuangan dan operasional otomatis, akurat, dan real-time\n"
    "✓ Jangkauan pasar lebih luas melalui platform e-commerce online\n"
    "✓ Kualitas pelayanan kepada pelanggan meningkat secara signifikan",
    fill_hex='FFF2CC', bold=False, center=False, font_size=10, width_cm=17.0
)

diagram_spacer()

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# BAB III – METODOLOGI PENELITIAN
# ════════════════════════════════════════════════════════════════════════════
heading1("BAB III\nMETODOLOGI PENELITIAN")
add_separator()

heading2("3.1 Jenis Penelitian")
body("Penelitian ini merupakan penelitian terapan (applied research) dengan metode penelitian dan pengembangan (Research and Development / R&D) yang berfokus pada perancangan dan pembangunan sistem informasi berbasis web. Pendekatan yang digunakan adalah pendekatan kualitatif untuk analisis kebutuhan sistem dan pendekatan kuantitatif untuk pengujian sistem menggunakan Black Box Testing.")
body('Menurut Sugiyono (2019), penelitian dan pengembangan adalah metode penelitian yang digunakan untuk menghasilkan produk tertentu dan menguji keefektifan produk tersebut. Dalam konteks ini, "produk" yang dimaksud adalah sistem informasi manajemen layanan perawatan kucing berbasis web untuk Cikal Pet Care Polewali Mandar.')

heading2("3.2 Lokasi dan Waktu Penelitian")

heading3("3.2.1 Lokasi Penelitian")
body("Penelitian ini dilaksanakan di Cikal Pet Care, yang berlokasi di Kabupaten Polewali Mandar, Provinsi Sulawesi Barat. Pemilihan lokasi ini didasarkan pada pertimbangan bahwa Cikal Pet Care merupakan salah satu usaha layanan perawatan kucing yang berkembang di wilayah tersebut namun belum memiliki sistem informasi manajemen yang memadai.")

heading3("3.2.2 Waktu Penelitian")
body("Penelitian ini dilaksanakan selama ± 4 bulan (Februari 2026 – Mei 2026) dengan rincian jadwal kegiatan sebagai berikut:")
add_table(
    headers=["No", "Kegiatan", "Feb 2026", "Mar 2026", "Apr 2026", "Mei 2026"],
    rows=[
        ["1", "Studi Literatur & Pengumpulan Data", "████", "", "", ""],
        ["2", "Analisis Kebutuhan Sistem", "████", "██", "", ""],
        ["3", "Perancangan Sistem (Design)", "", "████", "", ""],
        ["4", "Implementasi (Coding)", "", "██", "████", ""],
        ["5", "Pengujian Sistem (Testing)", "", "", "██", "██"],
        ["6", "Evaluasi & Perbaikan", "", "", "", "████"],
        ["7", "Penulisan Laporan", "██", "██", "██", "████"],
    ],
    col_widths=[0.8, 5.5, 2.5, 2.5, 2.5, 2.5]
)

heading2("3.3 Metode Pengumpulan Data")
body("Dalam penelitian ini, pengumpulan data dilakukan melalui beberapa teknik, yaitu:")

heading3("3.3.1 Observasi")
body("Observasi dilakukan dengan melakukan pengamatan langsung terhadap proses operasional bisnis di Cikal Pet Care Polewali Mandar. Pengamatan mencakup: proses pencatatan data pelanggan dan kucing, proses penerimaan dan pengelolaan booking layanan, proses pengelolaan stok produk, proses pembuatan laporan harian/bulanan, serta alur pembayaran yang berlaku.")

heading3("3.3.2 Wawancara")
body("Wawancara dilakukan secara terstruktur (structured interview) dengan narasumber kunci, yaitu: (1) Pemilik Cikal Pet Care — untuk mendapatkan gambaran umum bisnis, permasalahan yang dihadapi, dan kebutuhan sistem yang diinginkan; dan (2) Karyawan/Staff — untuk memahami alur kerja operasional sehari-hari dan kebutuhan sistem dari perspektif pengguna internal. Instrumen wawancara berupa daftar pertanyaan yang telah disiapkan sebelumnya dan dapat dikembangkan (semi-structured) mengikuti alur percakapan.")

heading3("3.3.3 Studi Dokumentasi")
body("Studi dokumentasi dilakukan dengan mengumpulkan dan menganalisis dokumen-dokumen yang berkaitan dengan operasional Cikal Pet Care, meliputi: catatan data pelanggan yang ada, catatan pesanan dan transaksi, daftar layanan dan harga yang berlaku, serta katalog produk yang dijual.")

heading3("3.3.4 Studi Literatur")
body("Studi literatur dilakukan dengan mengkaji referensi-referensi ilmiah yang relevan, meliputi: buku teks tentang rekayasa perangkat lunak, sistem informasi manajemen, dan basis data; jurnal ilmiah tentang pengembangan sistem informasi berbasis web dan layanan perawatan hewan; dokumentasi resmi teknologi yang digunakan; serta penelitian-penelitian terdahulu yang relevan.")

heading2("3.4 Metode Pengembangan Sistem")
body("Metode pengembangan sistem yang digunakan dalam penelitian ini adalah SDLC (Systems Development Life Cycle) dengan model Waterfall. Model Waterfall dipilih karena: (1) kebutuhan sistem telah terdefinisi dengan baik melalui hasil analisis dan wawancara; (2) cocok untuk proyek dengan ruang lingkup yang jelas dan tidak terlalu besar; (3) memberikan dokumentasi yang lengkap di setiap tahap; dan (4) mudah dikelola karena setiap fase memiliki milestone dan deliverable yang jelas (Pressman & Maxim, 2019).")

heading3("3.4.1 Tahap 1: Requirements Analysis (Analisis Kebutuhan)")
body("Pada tahap ini dilakukan identifikasi dan dokumentasi semua kebutuhan sistem berdasarkan hasil observasi dan wawancara. Kebutuhan dibagi menjadi dua kategori: kebutuhan fungsional dan kebutuhan non-fungsional.")

heading4("a. Kebutuhan Fungsional")
add_table(
    headers=["Kode", "Kebutuhan Fungsional", "Aktor"],
    rows=[
        ["KF-01", "Sistem menyediakan halaman beranda (homepage) yang menampilkan informasi umum tentang Cikal Pet Care", "Pelanggan"],
        ["KF-02", "Sistem menyediakan katalog layanan perawatan kucing yang dapat dilihat oleh pelanggan", "Pelanggan"],
        ["KF-03", "Sistem menyediakan formulir booking layanan online", "Pelanggan"],
        ["KF-04", "Sistem menyediakan katalog produk untuk dijual secara online", "Pelanggan"],
        ["KF-05", "Sistem menyediakan keranjang belanja (shopping cart)", "Pelanggan"],
        ["KF-06", "Sistem menyediakan fitur checkout dan pemilihan metode pembayaran", "Pelanggan"],
        ["KF-07", "Sistem menyediakan formulir booking paket penitipan kucing", "Pelanggan"],
        ["KF-08", "Sistem menyediakan halaman blog/artikel edukasi tentang perawatan kucing", "Pelanggan"],
        ["KF-09", "Sistem menyediakan halaman informasi kontak dan cara pembayaran", "Pelanggan"],
        ["KF-10", "Sistem menyediakan fitur login untuk administrator", "Admin"],
        ["KF-11", "Administrator dapat mengelola (tambah, ubah, hapus, lihat) data produk", "Admin"],
        ["KF-12", "Administrator dapat mengelola varian produk", "Admin"],
        ["KF-13", "Administrator dapat mengelola data layanan perawatan", "Admin"],
        ["KF-14", "Administrator dapat mengelola data booking layanan", "Admin"],
        ["KF-15", "Administrator dapat mengelola data pesanan termasuk verifikasi pembayaran", "Admin"],
        ["KF-16", "Administrator dapat mengelola data penitipan kucing", "Admin"],
        ["KF-17", "Administrator dapat mengelola data pelanggan", "Admin"],
        ["KF-18", "Administrator dapat mengelola konten blog dan artikel", "Admin"],
        ["KF-19", "Administrator dapat melihat dashboard dengan statistik dan grafik penjualan", "Admin"],
        ["KF-20", "Administrator dapat menghasilkan laporan penjualan dan operasional", "Admin"],
        ["KF-21", "Sistem mengirimkan notifikasi email otomatis kepada pelanggan saat booking dan pesanan", "Sistem"],
        ["KF-22", "Administrator dapat mengelola pengaturan website (informasi perusahaan, sosial media, dll.)", "Admin"],
        ["KF-23", "Sistem mendukung upload gambar produk dan layanan ke Cloudinary", "Admin"],
        ["KF-24", "Sistem menyediakan manajemen pengguna dengan role (Super Admin, Admin, Staff)", "Super Admin"],
    ],
    col_widths=[1.5, 12.0, 3.8]
)

heading4("b. Kebutuhan Non-Fungsional")
add_table(
    headers=["Kode", "Aspek", "Kebutuhan Non-Fungsional"],
    rows=[
        ["KNF-01", "Performa", "Halaman web harus dapat dimuat dalam waktu < 3 detik pada koneksi broadband standar"],
        ["KNF-02", "Keamanan", "Password pengguna harus di-hash menggunakan bcrypt sebelum disimpan ke database"],
        ["KNF-03", "Keamanan", "Halaman administrasi harus dilindungi autentikasi (tidak dapat diakses tanpa login)"],
        ["KNF-04", "Keamanan", "Sistem harus memvalidasi semua input pengguna untuk mencegah SQL Injection dan XSS"],
        ["KNF-05", "Usability", "Antarmuka sistem harus responsif (mobile-friendly) dan dapat digunakan di berbagai perangkat"],
        ["KNF-06", "Reliability", "Sistem harus memiliki uptime minimal 99% selama jam operasional"],
        ["KNF-07", "Maintainability", "Kode program harus terstruktur, terdokumentasi, dan mudah dipelihara"],
        ["KNF-08", "Compatibility", "Sistem harus dapat berjalan pada browser modern (Chrome, Firefox, Safari, Edge)"],
        ["KNF-09", "Scalability", "Sistem harus dapat menangani minimal 100 pengguna bersamaan tanpa penurunan performa signifikan"],
        ["KNF-10", "Data Integrity", "Sistem harus menjaga konsistensi dan integritas data melalui validasi input dan transaksi database"],
    ],
    col_widths=[1.5, 3.0, 12.8]
)

heading3("3.4.2 Tahap 2: System Design (Perancangan Sistem)")

heading4("3.4.2.1 Perancangan Arsitektur Sistem")
body("Sistem dirancang menggunakan arsitektur full-stack monolith dengan Next.js App Router. Arsitektur ini memungkinkan frontend (React Server Components + Client Components) dan backend (API Routes) berada dalam satu codebase yang terintegrasi, menyederhanakan deployment dan pemeliharaan. Sistem terdiri dari empat lapisan utama, yaitu:")
numbered_item("1", "Browser (Chrome/Firefox/Safari/Edge) dengan React Components, Tailwind CSS, dan Zustand untuk state management.", bold_prefix="Client Layer: ")
numbered_item("2", "Next.js 16 App Router yang menangani halaman publik, halaman admin, dan API Routes, dilengkapi NextAuth.js v5 untuk autentikasi dan Zod untuk validasi input.", bold_prefix="Application Layer: ")
numbered_item("3", "File-file TypeScript yang mengenkapsulasi logika bisnis: productService.ts, bookingService.ts, orderService.ts, emailService.ts (Resend API), dan settingsService.ts.", bold_prefix="Service Layer: ")
numbered_item("4", "Prisma ORM Client yang berinteraksi dengan database SQLite (better-sqlite3) untuk penyimpanan data, serta Cloudinary untuk penyimpanan media.", bold_prefix="Data Layer: ")

heading4("3.4.2.2 Perancangan Use Case Diagram")
body("Use case diagram sistem dibagi menjadi dua kelompok berdasarkan aktor yang terlibat:")

heading4("Use Case – Pelanggan (Public User)")
body("Aktor: Pelanggan (pengunjung website). Use Cases: UC-01 Melihat Halaman Beranda, UC-02 Melihat Katalog Layanan, UC-03 Melakukan Booking Layanan, UC-04 Melihat Katalog Produk, UC-05 Menambahkan Produk ke Keranjang, UC-06 Melakukan Checkout & Pembayaran, UC-07 Melihat Paket Penitipan, UC-08 Melakukan Booking Penitipan, UC-09 Melacak Status Pesanan, UC-10 Membaca Blog/Artikel, UC-11 Melihat Informasi Kontak.")

heading4("Use Case – Administrator")
body("Aktor: Administrator (Super Admin, Admin, Staff). Use Cases: UC-12 Login ke Sistem Admin, UC-13 Melihat Dashboard Analitik, UC-14 Mengelola Data Produk (CRUD), UC-15 Mengelola Varian Produk, UC-16 Mengelola Data Layanan (CRUD), UC-17 Mengelola Booking Layanan, UC-18 Mengelola Data Pesanan, UC-19 Verifikasi Pembayaran, UC-20 Mengelola Penitipan Kucing, UC-21 Mengelola Paket Penitipan, UC-22 Mengelola Data Pelanggan, UC-23 Mengelola Konten Blog, UC-24 Melihat Laporan Penjualan, UC-25 Mengatur Konfigurasi Website, UC-26 Mengelola Pengguna Admin, UC-27 Logout dari Sistem.")

heading4("3.4.2.3 Perancangan Activity Diagram")
body("Activity diagram dirancang untuk tiga proses utama sistem, yaitu proses booking layanan online oleh pelanggan, proses pembelian produk (checkout), dan proses login administrator. Setiap diagram menggambarkan alur kerja dari awal hingga akhir termasuk titik-titik pengambilan keputusan (decision nodes) dan kondisi alternatif yang mungkin terjadi.")
body("Pada proses booking layanan: pelanggan memilih layanan → sistem mengecek ketersediaan jadwal → pelanggan mengisi form → sistem memvalidasi input → sistem menyimpan data booking ke database → sistem mengirim email konfirmasi via Resend → tampilkan halaman konfirmasi. Jika jadwal tidak tersedia atau input tidak valid, sistem menampilkan pesan error yang sesuai.")
body("Pada proses pembelian produk: pelanggan memilih produk dan varian → tambah ke keranjang (Zustand store) → isi form checkout → pilih metode pembayaran → sistem memvalidasi stok → buat order di database → update stok → kirim email konfirmasi → tampilkan halaman sukses pesanan → pelanggan upload bukti pembayaran → admin verifikasi → update status order → notifikasi email ke pelanggan.")

heading4("3.4.2.4 Perancangan Sequence Diagram")
body("Sequence diagram dirancang untuk empat skenario utama sistem:")
numbered_item("1", "Menggambarkan interaksi dari pelanggan membuka halaman booking, memilih layanan, mengisi form, submit, sistem memvalidasi (Zod), menyimpan ke database melalui bookingService, dan mengirim email konfirmasi melalui Resend API.", bold_prefix="Sequence Diagram Booking Layanan Online: ")
numbered_item("2", "Menggambarkan alur dari pemilihan produk, penambahan ke keranjang (Zustand), checkout, validasi stok, pembuatan order dan order items di database, pengurangan stok, dan pengiriman email konfirmasi pesanan.", bold_prefix="Sequence Diagram Pemesanan Produk (Checkout): ")
numbered_item("3", "Menggambarkan alur admin mengakses daftar pesanan pending, mengklik tombol verifikasi, sistem mengupdate payment_status ke PAID dan status ke PROCESSING, mencatat ActivityLog, dan mengirim email notifikasi ke pelanggan.", bold_prefix="Sequence Diagram Verifikasi Pembayaran oleh Admin: ")
numbered_item("4", "Menggambarkan alur dari pembukaan halaman login, pengecekan sesi NextAuth, input kredensial, query database untuk mencari user, verifikasi password dengan bcrypt, pembuatan JWT session, dan redirect ke dashboard.", bold_prefix="Sequence Diagram Login Administrator: ")

heading4("3.4.2.5 Perancangan Entity Relationship Diagram (ERD)")
body("Sistem Cikal Pet Care memiliki 14 entitas utama dalam basis data dengan relasi-relasi sebagai berikut:")
add_table(
    headers=["Entitas 1", "Relasi", "Entitas 2", "Keterangan"],
    rows=[
        ["Product", "1 : N", "ProductVariant", "Satu produk memiliki banyak varian"],
        ["Product", "1 : N", "ProductReview", "Satu produk memiliki banyak ulasan"],
        ["Product", "1 : N", "OrderItem", "Satu produk dapat ada di banyak item pesanan"],
        ["ProductVariant", "1 : N", "OrderItem", "Satu varian dapat ada di banyak item pesanan"],
        ["Service", "1 : N", "ServiceBooking", "Satu layanan dapat dibooking berkali-kali"],
        ["Service", "1 : N", "OrderItem", "Satu layanan dapat ada di banyak item pesanan"],
        ["Customer", "1 : N", "Order", "Satu pelanggan dapat memiliki banyak pesanan"],
        ["Customer", "1 : N", "ServiceBooking", "Satu pelanggan dapat membuat banyak booking"],
        ["Customer", "1 : N", "PenitipanBooking", "Satu pelanggan dapat membuat banyak booking penitipan"],
        ["Order", "1 : N", "OrderItem", "Satu pesanan memiliki banyak item"],
        ["Order", "1 : N", "ActivityLog", "Satu pesanan dapat memiliki banyak log aktivitas"],
        ["PenitipanPackage", "1 : N", "PenitipanBooking", "Satu paket penitipan dapat dibooking berkali-kali"],
        ["User", "1 : N", "ActivityLog", "Satu pengguna admin dapat menghasilkan banyak log"],
    ],
    col_widths=[3.5, 2.0, 4.0, 7.8]
)

heading4("3.4.2.6 Perancangan Struktur Database")
body("Berikut adalah struktur tabel-tabel utama yang diimplementasikan dalam database SQLite menggunakan Prisma ORM:")

heading4("Tabel: products")
add_table(
    headers=["Nama Kolom", "Tipe Data", "Keterangan"],
    rows=[
        ["id", "TEXT (UUID)", "Primary Key"],
        ["name", "TEXT", "Nama produk"],
        ["slug", "TEXT UNIQUE", "Slug URL unik"],
        ["description", "TEXT", "Deskripsi produk"],
        ["sku", "TEXT UNIQUE", "Kode SKU unik"],
        ["price", "REAL", "Harga produk"],
        ["stock", "INTEGER", "Jumlah stok tersedia"],
        ["category", "TEXT", "Kategori produk"],
        ["image_url", "TEXT", "URL gambar utama (Cloudinary)"],
        ["gallery", "TEXT", "JSON array URL gambar tambahan"],
        ["is_active", "BOOLEAN", "Status aktif produk"],
        ["featured", "BOOLEAN", "Produk unggulan/featured"],
        ["low_stock_alert", "INTEGER", "Batas minimum stok untuk peringatan"],
        ["created_at", "DATETIME", "Waktu data dibuat"],
        ["updated_at", "DATETIME", "Waktu data terakhir diperbarui"],
    ],
    col_widths=[4.0, 4.0, 9.3]
)

heading4("Tabel: service_bookings")
add_table(
    headers=["Nama Kolom", "Tipe Data", "Keterangan"],
    rows=[
        ["id", "TEXT (UUID)", "Primary Key"],
        ["service_id", "TEXT", "Foreign Key → services.id"],
        ["customer_id", "TEXT", "Foreign Key → customers.id"],
        ["booking_date", "DATETIME", "Tanggal booking layanan"],
        ["booking_time", "TEXT", "Jam booking layanan"],
        ["pet_name", "TEXT", "Nama kucing yang dibooking"],
        ["pet_type", "TEXT", "Jenis/ras kucing"],
        ["notes", "TEXT", "Catatan tambahan dari pelanggan"],
        ["status", "TEXT", "Status: PENDING, CONFIRMED, COMPLETED, CANCELED"],
        ["created_at", "DATETIME", "Waktu data dibuat"],
        ["updated_at", "DATETIME", "Waktu data terakhir diperbarui"],
    ],
    col_widths=[4.0, 4.0, 9.3]
)

heading4("Tabel: orders")
add_table(
    headers=["Nama Kolom", "Tipe Data", "Keterangan"],
    rows=[
        ["id", "TEXT (UUID)", "Primary Key"],
        ["order_number", "TEXT UNIQUE", "Nomor pesanan unik"],
        ["customer_id", "TEXT", "Foreign Key → customers.id"],
        ["subtotal", "REAL", "Subtotal sebelum ongkir & diskon"],
        ["shipping_cost", "REAL", "Biaya pengiriman"],
        ["discount_amount", "REAL", "Jumlah diskon"],
        ["total_amount", "REAL", "Total yang harus dibayar"],
        ["payment_method", "TEXT", "Metode pembayaran (BANK_TRANSFER, QRIS, COD)"],
        ["payment_status", "TEXT", "Status pembayaran (PENDING, VERIFYING, PAID, FAILED)"],
        ["payment_proof_url", "TEXT", "URL gambar bukti transfer (Cloudinary)"],
        ["status", "TEXT", "Status pesanan (PENDING, PROCESSING, SHIPPED, COMPLETED)"],
        ["shipping_address", "TEXT", "Alamat pengiriman"],
        ["tracking_number", "TEXT", "Nomor resi pengiriman"],
        ["created_at", "DATETIME", "Waktu data dibuat"],
        ["updated_at", "DATETIME", "Waktu data terakhir diperbarui"],
    ],
    col_widths=[4.0, 4.0, 9.3]
)

heading4("3.4.2.7 Perancangan Antarmuka (UI/UX Design)")
body("Antarmuka sistem dirancang mengikuti prinsip mobile-first dengan warna utama biru (#2563EB) yang merepresentasikan kepercayaan dan profesionalisme. Perancangan dibagi menjadi dua kelompok halaman:")
body("Halaman publik yang dirancang meliputi: (1) Halaman Beranda — hero section, statistik layanan, daftar layanan unggulan, produk featured, testimoni, dan blog terbaru; (2) Halaman Layanan — grid kartu layanan dengan harga, durasi, dan tombol booking; (3) Halaman Booking — form booking dengan validasi real-time; (4) Halaman Produk — grid produk dengan filter kategori dan keranjang belanja; (5) Halaman Checkout — form data pengiriman, ringkasan pesanan, dan pemilihan pembayaran; (6) Halaman Blog — grid artikel dengan filter dan pencarian; serta halaman-halaman pendukung lainnya.")
body("Panel administrasi menggunakan layout sidebar dengan: Sidebar kiri — menu navigasi antar modul; Header atas — info pengguna yang login dan tombol logout; serta konten utama di sisi kanan. Halaman admin yang dirancang mencakup: Dashboard (statistik dan grafik), Manajemen Produk, Manajemen Layanan, Manajemen Booking, Manajemen Pesanan, Manajemen Penitipan, Manajemen Blog, Laporan, dan Pengaturan Website.")

heading3("3.4.3 Tahap 3: Implementation (Implementasi)")
body("Pada tahap implementasi, rancangan sistem diwujudkan dalam bentuk kode program menggunakan Next.js 16 App Router dengan TypeScript. Implementasi dilakukan mengikuti struktur folder Next.js App Router dengan pemisahan yang jelas antara halaman publik (/app/(public)/), halaman admin (/app/admin/), endpoint API (/app/api/), komponen UI reusable (/src/components/), service layer (/src/services/), dan schema database (/prisma/schema.prisma).")
body("Implementasi keamanan sistem meliputi: (1) Autentikasi berbasis JWT menggunakan NextAuth.js v5 dengan Credentials Provider; (2) Middleware Next.js yang memproteksi semua route /admin/* dan mengecek validitas sesi; (3) Password hashing menggunakan bcryptjs dengan salt round 12; (4) Validasi input menggunakan Zod schema di setiap endpoint API untuk mencegah data tidak valid; (5) Parameterized queries otomatis oleh Prisma ORM untuk mencegah SQL Injection; dan (6) Penggunaan HTTPS pada environment produksi.")

heading3("3.4.4 Tahap 4: Testing (Pengujian Sistem)")
body("Pengujian sistem dilakukan menggunakan metode Black Box Testing yang berfokus pada pengujian fungsionalitas eksternal sistem. Berikut adalah rencana pengujian (test cases) yang digunakan:")
add_table(
    headers=["No", "Modul", "Skenario Pengujian", "Input", "Output yang Diharapkan"],
    rows=[
        ["TC-01", "Login Admin", "Login dengan kredensial valid", "Email & password benar", "Redirect ke /admin/dashboard"],
        ["TC-02", "Login Admin", "Login dengan password salah", "Email benar, password salah", "Pesan error 'Email atau password salah'"],
        ["TC-03", "Login Admin", "Login dengan email tidak terdaftar", "Email tidak terdaftar", "Pesan error 'Email atau password salah'"],
        ["TC-04", "Login Admin", "Akses halaman admin tanpa login", "URL /admin/dashboard langsung", "Redirect otomatis ke /admin/login"],
        ["TC-05", "Produk", "Tambah produk baru dengan data lengkap", "Data produk valid lengkap", "Produk tersimpan dan muncul di tabel"],
        ["TC-06", "Produk", "Tambah produk dengan SKU duplikat", "SKU yang sudah ada di database", "Pesan error 'SKU sudah digunakan'"],
        ["TC-07", "Produk", "Edit data produk yang ada", "Data produk diubah (nama, harga)", "Data produk ter-update di database"],
        ["TC-08", "Produk", "Hapus produk", "Klik tombol hapus produk", "Produk terhapus dari database dan tabel"],
        ["TC-09", "Booking", "Booking layanan dengan data lengkap", "Form booking diisi lengkap", "Booking tersimpan, email konfirmasi terkirim"],
        ["TC-10", "Booking", "Booking dengan tanggal yang sudah lewat", "Tanggal masa lalu", "Pesan error validasi tanggal"],
        ["TC-11", "Pesanan", "Buat pesanan dengan produk valid", "Data checkout dan pembayaran lengkap", "Pesanan tersimpan, email konfirmasi terkirim"],
        ["TC-12", "Pesanan", "Upload bukti pembayaran", "File gambar (JPG/PNG)", "Gambar ter-upload ke Cloudinary"],
        ["TC-13", "Pesanan", "Verifikasi pembayaran oleh admin", "Admin konfirmasi pembayaran", "Status berubah ke PAID, email notifikasi terkirim"],
        ["TC-14", "Dashboard", "Lihat halaman dashboard admin", "Admin login dan akses dashboard", "Statistik dan grafik tampil dengan benar"],
        ["TC-15", "Laporan", "Generate laporan penjualan", "Pilih periode laporan", "Laporan tampil dengan data yang akurat"],
        ["TC-16", "Penitipan", "Booking paket penitipan kucing", "Form penitipan diisi lengkap", "Booking tersimpan, konfirmasi tampil"],
        ["TC-17", "Blog", "Tambah artikel blog baru", "Konten artikel diisi lengkap", "Artikel tersimpan dan dapat dilihat di frontend"],
        ["TC-18", "Pengaturan", "Ubah informasi kontak perusahaan", "Data kontak baru di-input", "Pengaturan ter-update dan tampil di website"],
        ["TC-19", "Keranjang", "Tambah produk ke keranjang belanja", "Pilih produk dan jumlah", "Item muncul di keranjang, total dihitung benar"],
        ["TC-20", "Keranjang", "Checkout dengan keranjang kosong", "Akses checkout tanpa produk", "Pesan 'Keranjang belanja kosong'"],
    ],
    col_widths=[0.7, 2.2, 3.5, 3.5, 7.4]
)

heading3("3.4.5 Tahap 5: Deployment (Penerapan)")
body("Pada tahap ini, sistem yang telah diuji dan dinyatakan layak akan diterapkan pada lingkungan produksi. Tahapan deployment meliputi: (1) Persiapan server/hosting dan konfigurasi lingkungan; (2) Konfigurasi environment variables (DATABASE_URL, NEXTAUTH_SECRET, CLOUDINARY_API_KEY, RESEND_API_KEY); (3) Migrasi database dengan menjalankan Prisma migrate; (4) Build aplikasi dengan perintah next build; (5) Seeding data awal menggunakan Prisma seed; (6) Konfigurasi domain dan SSL/HTTPS certificate; serta (7) Pelatihan administrator dan staff Cikal Pet Care dalam menggunakan sistem.")

heading2("3.5 Definisi Operasional")
body("Untuk menyamakan persepsi dan menghindari kesalahan penafsiran terhadap variabel dan konsep yang digunakan dalam penelitian ini, berikut disajikan definisi operasional dari istilah-istilah kunci:")
add_table(
    headers=["No", "Istilah", "Definisi Operasional"],
    rows=[
        ["1", "Sistem Informasi Manajemen (SIM)", "Aplikasi web berbasis Next.js yang dibangun untuk mengelola seluruh proses operasional Cikal Pet Care, meliputi modul produk, layanan, booking, penitipan, pesanan, pelanggan, laporan, blog, dan pengaturan website"],
        ["2", "Layanan Perawatan Kucing", "Jasa yang disediakan Cikal Pet Care meliputi grooming, konsultasi kesehatan, vaksinasi, dan layanan lain yang terdaftar pada tabel services dalam database"],
        ["3", "Booking Layanan", "Proses pemesanan jadwal layanan perawatan oleh pelanggan melalui formulir online yang menghasilkan record baru pada tabel service_bookings dengan status awal PENDING"],
        ["4", "Penitipan Kucing (Cat Boarding)", "Layanan akomodasi sementara untuk kucing yang dikelola melalui modul penitipan, di mana pelanggan memilih paket, tanggal check-in dan check-out, tersimpan dalam tabel penitipan_bookings"],
        ["5", "Pesanan (Order)", "Transaksi pembelian produk oleh pelanggan melalui fitur e-commerce, yang mencakup proses keranjang belanja, checkout, dan pembayaran, tersimpan dalam tabel orders dan order_items"],
        ["6", "Verifikasi Pembayaran", "Proses konfirmasi pembayaran oleh administrator setelah pelanggan mengunggah bukti transfer, mengubah payment_status dari VERIFYING menjadi PAID"],
        ["7", "Administrator", "Pengguna yang memiliki akun login di panel admin dengan role SUPER_ADMIN, ADMIN, atau STAFF, tersimpan dalam tabel users"],
        ["8", "Pelanggan (Customer)", "Pengguna yang menggunakan layanan atau membeli produk Cikal Pet Care tanpa akun login; diidentifikasi unik berdasarkan nomor telepon dalam tabel customers"],
        ["9", "Dashboard Analitik", "Halaman /admin yang menampilkan ringkasan statistik bisnis beserta grafik tren menggunakan Recharts, di-generate dari query database secara real-time"],
        ["10", "Notifikasi Email Otomatis", "Pesan elektronik yang dikirim secara otomatis oleh sistem melalui Resend API pada event tertentu tanpa intervensi manual administrator"],
        ["11", "Grooming", "Layanan perawatan fisik kucing mencakup mandi, pengeringan, penyisiran, trimming, potong kuku, dan pembersihan telinga/mata"],
        ["12", "Paket Penitipan", "Pilihan layanan penitipan dengan fasilitas dan harga per malam yang berbeda, dikelola dalam tabel penitipan_packages"],
        ["13", "Black Box Testing", "Metode pengujian yang memverifikasi fungsionalitas sistem dengan membandingkan output aktual dengan output yang diharapkan tanpa melihat kode internal"],
        ["14", "Slug", "String URL-friendly unik untuk setiap produk, layanan, dan artikel blog (contoh: paket-grooming-premium), tersimpan dalam kolom slug tiap tabel"],
        ["15", "SKU", "Kode unik yang mengidentifikasi setiap produk atau varian produk dalam sistem inventori, tersimpan dalam kolom sku pada tabel products dan product_variants"],
    ],
    col_widths=[0.7, 3.5, 13.1]
)

heading2("3.6 Alat dan Bahan Penelitian")

heading3("3.6.1 Perangkat Keras (Hardware)")
add_table(
    headers=["Komponen", "Spesifikasi Minimum"],
    rows=[
        ["Laptop/PC", "Prosesor Intel Core i5/i7 atau AMD Ryzen 5/7"],
        ["RAM", "8 GB (minimum), 16 GB (disarankan)"],
        ["Storage", "SSD 256 GB (minimum)"],
        ["Koneksi Internet", "Broadband minimum 10 Mbps"],
    ],
    col_widths=[5.0, 12.3]
)

heading3("3.6.2 Perangkat Lunak (Software)")
add_table(
    headers=["Perangkat Lunak", "Versi", "Fungsi"],
    rows=[
        ["Node.js", "v18.x / v20.x", "Runtime JavaScript server-side"],
        ["Next.js", "v16.1.6", "Framework web full-stack"],
        ["React", "v18.3.1", "Library UI"],
        ["TypeScript", "v5.x", "Bahasa pemrograman (superset JavaScript)"],
        ["Prisma ORM", "v5.22.0", "Database Object Relational Mapper"],
        ["SQLite (better-sqlite3)", "v12.6.2", "Database relasional serverless"],
        ["Tailwind CSS", "v3.x", "CSS Framework utility-first"],
        ["NextAuth.js", "v5.0.0-beta", "Library autentikasi"],
        ["Cloudinary", "v2.9.0", "Platform penyimpanan & CDN media cloud"],
        ["Resend", "v6.9.2", "Layanan pengiriman email transaksional"],
        ["Zustand", "v5.0.11", "State management library"],
        ["Recharts", "v3.8.1", "Library visualisasi data / grafik"],
        ["Zod", "v4.3.6", "Library validasi schema TypeScript"],
        ["Bcryptjs", "v3.0.3", "Library hashing password"],
        ["VS Code", "Terbaru", "Code editor utama"],
        ["Git", "Terbaru", "Version control system"],
        ["Postman", "Terbaru", "API testing tool"],
        ["Figma", "Berbasis web", "Tool desain UI/UX"],
    ],
    col_widths=[4.0, 3.5, 9.8]
)

heading2("3.7 Diagram Alir Penelitian")
body("Diagram alir (flowchart) berikut menggambarkan tahapan-tahapan penelitian secara sistematis, mulai dari identifikasi masalah hingga penyelesaian penulisan laporan skripsi. Simbol persegi panjang (  ) menyatakan proses, simbol belah ketupat (◇) menyatakan keputusan, dan simbol oval menyatakan terminal (mulai/selesai).")

# ── Flowchart Penelitian BAB 3 ─────────────────────────────────────────────

# MULAI
diagram_box("● MULAI", fill_hex='1F4E79', bold=True,
            font_size=12, width_cm=7.0, border_hex='1F4E79')
diagram_arrow()

# Step 1
diagram_box(
    "1. IDENTIFIKASI MASALAH\n"
    "Observasi langsung & wawancara di Cikal Pet Care Polewali Mandar\n"
    "→ Dokumentasi permasalahan operasional yang ada",
    fill_hex='DEEAF1', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Step 2
diagram_box(
    "2. STUDI LITERATUR & PENGUMPULAN DATA\n"
    "• Kajian buku teks: SI, rekayasa PL, basis data, metode penelitian\n"
    "• Kajian jurnal ilmiah: penelitian terdahulu yang relevan\n"
    "• Kajian dokumentasi teknologi: Next.js, Prisma, TypeScript, dll.\n"
    "• Studi dokumentasi: catatan operasional Cikal Pet Care",
    fill_hex='DEEAF1', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Decision 1
diagram_box(
    "◇  APAKAH DATA YANG TERKUMPUL SUDAH CUKUP?",
    fill_hex='FFF2CC', bold=True, center=True, font_size=10, width_cm=17.0,
    border_hex='D4A017'
)
diagram_arrow("↓  YA")

# Step 3
diagram_box(
    "3. ANALISIS KEBUTUHAN SISTEM\n"
    "• Kebutuhan Fungsional (KF-01 s.d. KF-24)\n"
    "• Kebutuhan Non-Fungsional (KNF-01 s.d. KNF-10)\n"
    "• Penetapan ruang lingkup & batasan sistem",
    fill_hex='E2EFDA', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Step 4
diagram_box(
    "4. PERANCANGAN SISTEM\n"
    "• Arsitektur sistem (full-stack Next.js App Router)\n"
    "• Use Case Diagram (pelanggan & administrator)\n"
    "• Activity Diagram (booking, pemesanan, login)\n"
    "• Sequence Diagram (4 proses utama)\n"
    "• Entity Relationship Diagram (ERD — 14 entitas)\n"
    "• Perancangan struktur database (Prisma schema)\n"
    "• Perancangan antarmuka / UI-UX (Figma)",
    fill_hex='E2EFDA', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Step 5
diagram_box(
    "5. IMPLEMENTASI / CODING\n"
    "• Pengembangan backend: Next.js API Routes + Prisma ORM + SQLite\n"
    "• Pengembangan frontend: React Server Components + Tailwind CSS\n"
    "• Integrasi autentikasi: NextAuth.js v5 (Credentials + JWT)\n"
    "• Integrasi media: Cloudinary (upload & delivery gambar)\n"
    "• Integrasi email: Resend API (notifikasi otomatis)\n"
    "• State management: Zustand (keranjang belanja)\n"
    "• Validasi data: Zod schema validation",
    fill_hex='EDE7F6', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Step 6
diagram_box(
    "6. PENGUJIAN SISTEM\n"
    "Metode: Black Box Testing (20 skenario uji coba)\n"
    "Cakupan: Login, Produk, Booking, Pesanan, Penitipan,\n"
    "Dashboard, Laporan, Blog, Keranjang, Upload Gambar",
    fill_hex='EDE7F6', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Decision 2
diagram_box(
    "◇  APAKAH SISTEM LOLOS SEMUA SKENARIO PENGUJIAN?",
    fill_hex='FFF2CC', bold=True, center=True, font_size=10, width_cm=17.0,
    border_hex='D4A017'
)
# TIDAK path note
p_no = doc.add_paragraph()
p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_no.paragraph_format.space_before = Pt(0)
p_no.paragraph_format.space_after = Pt(0)
r_no = p_no.add_run("↙ TIDAK: Perbaikan bug & kembali ke tahap Implementasi")
set_font(r_no, 10, italic=True)
r_no.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

diagram_arrow("↓  YA")

# Step 7
diagram_box(
    "7. DEPLOYMENT & PENERAPAN SISTEM\n"
    "• Konfigurasi server/hosting & environment variables\n"
    "• Migrasi database (prisma migrate deploy)\n"
    "• Build aplikasi (next build) & optimasi produksi\n"
    "• Seeding data awal (admin, layanan, paket penitipan)\n"
    "• Konfigurasi domain & SSL/HTTPS certificate\n"
    "• Pelatihan pengguna (administrator & staff Cikal Pet Care)",
    fill_hex='FFE8E8', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# Step 8
diagram_box(
    "8. PENULISAN LAPORAN SKRIPSI\n"
    "• Dokumentasi seluruh tahapan penelitian\n"
    "• Analisis & pembahasan hasil penelitian\n"
    "• Penarikan kesimpulan dan saran pengembangan",
    fill_hex='FFE8E8', bold=False, center=False, font_size=10, width_cm=17.0
)
diagram_arrow()

# SELESAI
diagram_box("● SELESAI", fill_hex='1F4E79', bold=True,
            font_size=12, width_cm=7.0, border_hex='1F4E79')

diagram_spacer()

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# DAFTAR PUSTAKA
# ════════════════════════════════════════════════════════════════════════════
heading1("DAFTAR PUSTAKA")
add_separator()

def add_ref_para(doc, ref_text):
    """Tambah paragraf referensi APA dengan italic otomatis via marker [i]...[/i]."""
    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=4,
                    line_spacing=18,
                    left_indent=1.27, first_line_indent=-1.27)
    segments = []
    remaining = ref_text
    while '[i]' in remaining:
        before = remaining[:remaining.index('[i]')]
        remaining = remaining[remaining.index('[i]') + 3:]
        end_idx = remaining.index('[/i]')
        italic_text = remaining[:end_idx]
        remaining = remaining[end_idx + 4:]
        if before:
            segments.append((before, False))
        segments.append((italic_text, True))
    if remaining:
        segments.append((remaining, False))
    for text, is_italic in segments:
        r = p.add_run(text)
        r.italic = is_italic
        set_font(r, 12)

refs = [
    "Ammann, P., & Offutt, J. (2022). [i]Introduction to Software Testing[/i] (2nd ed.). Cambridge University Press.",
    "American Association of Feline Practitioners (AAFP). (2021). [i]Feline Grooming and Coat Care Guidelines[/i]. American Association of Feline Practitioners. https://catvets.com/guidelines/practice-guidelines",
    "American Pet Products Association (APPA). (2023). [i]2023-2024 APPA National Pet Owners Survey[/i]. Greenwich: American Pet Products Association.",
    "Auth.js. (2024). [i]Auth.js Documentation - Authentication for the web[/i]. https://authjs.dev/",
    "Azzahra, F., Nurhidayat, M., & Ramadhan, A. (2023). Penerapan sistem e-commerce pada UMKM perlengkapan hewan peliharaan menggunakan ReactJS. [i]Jurnal Informatika dan Rekayasa Perangkat Lunak[/i], [i]5[/i](2), 112–124.",
    "better-sqlite3. (2024). [i]better-sqlite3: Documentation and API Reference[/i]. https://github.com/WiseLibs/better-sqlite3",
    "Campbell, R. (2021). [i]Pet Boarding and Grooming Industry: Market Analysis and Best Practices[/i]. Pet Industry Joint Advisory Council (PIJAC).",
    "Cloudinary. (2024). [i]Cloudinary Documentation: Image and Video Management Platform[/i]. https://cloudinary.com/documentation",
    "Colinhacks. (2024). [i]Zod Documentation: TypeScript-first Schema Validation[/i]. https://zod.dev/",
    "Dennis, A., Wixom, B. H., & Tegarden, D. (2021). [i]Systems Analysis and Design: An Object-Oriented Approach with UML[/i] (6th ed.). John Wiley & Sons.",
    "Drangies, D. (2024). [i]bcryptjs: Optimized bcrypt in JavaScript[/i]. https://github.com/dcodeIO/bcrypt.js",
    "Elmasri, R., & Navathe, S. B. (2023). [i]Fundamentals of Database Systems[/i] (8th ed.). Pearson Education.",
    "Evergreen, S. D. H. (2022). [i]Effective Data Visualization: The Right Chart for the Right Data[/i] (3rd ed.). SAGE Publications.",
    "Fitriani, S., & Rahmat, A. (2022). Analisis pengembangan bisnis perawatan kucing di Indonesia: Peluang dan tantangan di era digital. [i]Jurnal Manajemen dan Bisnis Indonesia[/i], [i]10[/i](1), 45–62.",
    "Gaur, A., & Sharma, R. (2022). Online appointment booking systems: A review of design principles and user experience factors. [i]International Journal of Computer Applications[/i], [i]184[/i](15), 1–8.",
    "Hipp, R. D. (2024). [i]SQLite Home Page: Small. Fast. Reliable. Choose any three[/i]. https://www.sqlite.org/",
    "International Boarding & Pet Services Association (IBPSA). (2022). [i]Standards of Care for Pet Boarding Facilities[/i]. International Boarding & Pet Services Association.",
    "Kadir, A. (2020). [i]Pengantar Sistem Informasi[/i] (Edisi Revisi). ANDI Offset.",
    "Kurniawan, R., Pratiwi, M., & Hidayat, T. (2022). Pengembangan sistem informasi pet boarding berbasis web dengan pendekatan prototyping. [i]Jurnal Teknologi dan Sistem Informasi[/i], [i]3[/i](2), 98–114.",
    "Laudon, K. C., & Laudon, J. P. (2020). [i]Management Information Systems: Managing the Digital Firm[/i] (16th ed.). Pearson Education.",
    "Mahendra, I. G. B., & Rai, I. G. A. (2022). Perancangan sistem informasi veteriner berbasis web dengan fitur notifikasi email otomatis. [i]Jurnal TEKNOIF Teknik Informatika[/i], [i]10[/i](2), 78–89.",
    "Meta Open Source. (2024). [i]React Documentation: The library for web and native user interfaces[/i]. https://react.dev/",
    "Microsoft. (2024). [i]TypeScript Documentation: JavaScript With Syntax For Types[/i]. https://www.typescriptlang.org/docs/",
    "Next.js Documentation. (2024). [i]Next.js by Vercel - The React Framework[/i]. https://nextjs.org/docs",
    "O'Brien, J. A., & Marakas, G. M. (2011). [i]Management Information Systems[/i] (10th ed.). McGraw-Hill Education.",
    "OWASP. (2024). [i]Password Storage Cheat Sheet[/i]. Open Web Application Security Project. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
    "Pmndrs. (2024). [i]Zustand Documentation: A small, fast and scalable bearbones state-management solution[/i]. https://zustand-demo.pmnd.rs/",
    "Populix. (2023). [i]Tren Kepemilikan Hewan Peliharaan di Indonesia 2023[/i]. Populix Research. https://populix.co/articles/tren-hewan-peliharaan-indonesia",
    "Pressman, R. S., & Maxim, B. R. (2019). [i]Software Engineering: A Practitioner's Approach[/i] (9th ed.). McGraw-Hill Education.",
    "Prisma. (2024). [i]Prisma Documentation: Next-generation Node.js and TypeScript ORM[/i]. https://www.prisma.io/docs",
    "Purnama, R., & Dewi, L. A. (2023). Sistem informasi penjualan dan penitipan hewan peliharaan berbasis web (Studi kasus: Rumah Anabul Yogyakarta). [i]Jurnal Informatika: Jurnal Pengembangan IT (JPIT)[/i], [i]8[/i](1), 34–43.",
    "Rahayu, D., Sari, R. P., & Nugroho, A. (2021). Sistem informasi manajemen pet shop berbasis web menggunakan framework Laravel. [i]Jurnal Pengembangan Teknologi Informasi dan Ilmu Komputer (J-PTIIK)[/i], [i]5[/i](7), 3210–3220.",
    "Rainer, R. K., & Prince, B. (2023). [i]Introduction to Information Systems[/i] (9th ed.). John Wiley & Sons.",
    "Recharts. (2024). [i]Recharts Documentation: Redefined chart library built with React and D3[/i]. https://recharts.org/",
    "Resend. (2024). [i]Resend Documentation: Email for developers[/i]. https://resend.com/docs",
    "Santoso, B., & Wibowo, A. H. (2020). Rancang bangun sistem informasi klinik veteriner berbasis web untuk meningkatkan kualitas pelayanan. [i]Seminar Nasional Teknologi Informasi dan Komunikasi (SENTIKA)[/i], [i]2020[/i](1), 45–54.",
    "Satzinger, J. W., Jackson, R. B., & Burd, S. D. (2022). [i]Systems Analysis and Design in a Changing World[/i] (8th ed.). Cengage Learning.",
    "Schlueter, C., & Grandin, T. (2019). [i]Feline Behavioral Health and Welfare[/i]. Iowa State University Press.",
    "Sommerville, I. (2021). [i]Engineering Software Products: An Introduction to Modern Software Engineering[/i]. Pearson Education.",
    "Sugiyono. (2019). [i]Metode Penelitian Kuantitatif, Kualitatif, dan R&D[/i] (Edisi ke-2). Alfabeta.",
    "Tailwind Labs. (2024). [i]Tailwind CSS Documentation: A utility-first CSS framework[/i]. https://tailwindcss.com/docs",
    "Turban, E., Pollard, C., & Wood, G. (2021). [i]Information Technology for Management: On-Demand Strategies for Performance, Growth and Sustainability[/i] (12th ed.). John Wiley & Sons.",
    "TypeScript Documentation. (2024). [i]TypeScript Handbook[/i]. https://www.typescriptlang.org/docs/handbook/intro.html",
    "Vercel. (2024). [i]Next.js Documentation: The React Framework for Production[/i]. https://nextjs.org",
    "Wijaya, H., & Susanto, E. (2023). Implementasi Next.js dan Prisma ORM dalam pengembangan sistem informasi manajemen berbasis web: Studi perbandingan performa. [i]Jurnal Sistem Informasi dan Teknik Komputer (JOSISTKOM)[/i], [i]7[/i](2), 156–170.",
    "Zustand Documentation. (2024). [i]Zustand: Bear necessities for state management in React[/i]. https://github.com/pmndrs/zustand",
]

for ref in refs:
    add_ref_para(doc, ref)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = r"e:\website kuicng adit\SKRIPSI_Cikal_Pet_Care.docx"
doc.save(output_path)
print(f"[OK] File Word berhasil dibuat: {output_path}")
